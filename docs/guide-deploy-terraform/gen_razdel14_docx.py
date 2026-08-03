# -*- coding: utf-8 -*-
"""Generate docs/Razdel-14-helm.docx — Section 14 Helm chart."""
from __future__ import annotations

import importlib.util
import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches

ROOT = Path(r"D:/Project_infra/greeting-service-infra")
TEMPLATE_CANDIDATES = [
    ROOT / "docs/Razdel-13-kubernetes-manual.docx",
    ROOT / "docs/Traffic-Visibility-Caretta.docx",
    ROOT / "docs/Razdel-flyway-migrations.docx",
    ROOT / "docs/Razdel-14-helm.docx",
]
OUTPUT = ROOT / "docs/Razdel-14-helm.docx"
BUILD = ROOT / "docs/_build_Razdel-14-helm.docx"
DIAGRAM_SCRIPT = ROOT / "docs/helm_diagram_render.py"
IMAGE_ECOSYSTEM = ROOT / "docs/images/helm/helm-kubernetes-ecosystem.png"
IMAGE_PROJECT = ROOT / "docs/images/helm/helm-flow.png"
INCIDENTA_IMAGE_URL = "https://app.incidenta.tech/static/images/articles/0031-tools-helm.png"

G11 = Path(r"C:/Users/sky/AppData/Local/Temp/gen_razdel11a_docx.py")
spec = importlib.util.spec_from_file_location("g11", G11)
g11 = importlib.util.module_from_spec(spec)
assert spec.loader
spec.loader.exec_module(g11)

HELM_CHARTS_URL = "https://helm.sh/docs/topics/charts/"
HELM_UPGRADE_URL = "https://helm.sh/docs/helm/helm_upgrade/"
INCIDENTA_HELM_URL = "https://app.incidenta.tech/article/tools-helm/"

CHART = "infra/helm/greeting-service"
RELEASE = "greeting-service"
NS_DEV = "dev"
NS_PROD = "prod"
DEVTOOLS_IP = "72.56.249.137"
HOST_DEV = "greeting-dev.cloud-terra.online"
KUBECONFIG_GB = "export KUBECONFIG=/c/Users/sky/.kube/timeweb-greeting.yaml"
REPO_CD = "cd '/d/Project_infra/greeting-service-infra'"

TOC = [
    ("14.1. Цель", "sec14_goal"),
    ("14.2. Предварительные условия", "sec14_prereq"),
    ("14.3. Теория: Helm и chart", "sec14_theory"),
    ("14.4. Схема работы Helm в экосистеме Kubernetes", "sec14_diagram"),
    ("14.5. Структура chart greeting-service", "sec14_structure"),
    ("14.6. Проверка chart: lint и template", "sec14_lint"),
    ("14.7. Проверка тега в registry и деплой в dev", "sec14_dev"),
    ("14.8. Деплой в prod", "sec14_prod"),
    ("14.9. Управление release", "sec14_manage"),
    ("14.10. Откат и удаление", "sec14_rollback"),
    ("14.11. Как проверить результат", "sec14_verify"),
    ("14.12. Типичные ошибки", "sec14_errors"),
]


def add_external_doc_citation(doc: Document, url: str, quote_en: str, quote_ru: str) -> None:
    g11.add_empty(doc)
    g11.add_normal(doc, f"Источник: {url}")
    g11.add_empty(doc)
    g11.add_normal(doc, "Цитата:")
    g11.add_platform_block(doc, "Оригинал (English)", quote_en)
    g11.add_empty(doc)
    g11.add_normal(doc, "Перевод:")
    g11.add_platform_block(doc, "Перевод на русский", quote_ru)


def add_toc(doc: Document) -> None:
    g11.add_heading2(doc, "Оглавление", "sec14_toc")
    for title, anchor in TOC:
        indent = bool(re.match(r"14\.\d+\.\d+", title))
        g11.add_hyperlink_paragraph(doc, title, anchor, indent=indent)
    g11.add_empty(doc)


def add_gitbash_block(doc: Document, code: str) -> None:
    g11.add_platform_block(
        doc,
        "Локальный ПК — Windows (Git Bash), корень репозитория",
        f"{REPO_CD}\n{KUBECONFIG_GB}\n{code}",
    )


def add_figure(doc: Document, path: Path, caption: str) -> None:
    g11.add_empty(doc)
    g11.add_normal(doc, caption, consolas=True)
    g11.add_empty(doc)
    doc.add_picture(str(path), width=Inches(6.5))
    g11.add_empty(doc)


def ensure_diagrams() -> None:
    fetch_script = ROOT / "docs/fetch_incidenta_helm_image.py"
    if not IMAGE_ECOSYSTEM.exists():
        subprocess.run([sys.executable, str(fetch_script)], check=True)
    if not IMAGE_PROJECT.exists():
        subprocess.run([sys.executable, str(DIAGRAM_SCRIPT)], check=True)


def resolve_template() -> Path:
    for path in TEMPLATE_CANDIDATES:
        if path.exists():
            return path
    raise FileNotFoundError("No Word template found in docs/")


def main() -> None:
    ensure_diagrams()
    g11._bookmark_id = 0
    template = resolve_template()
    shutil.copy2(template, BUILD)
    doc = Document(BUILD)
    g11.clear_document_body(doc)

    g11.add_normal(doc, "Раздел 14. Helm chart: деплой и управление приложением", consolas=True)
    g11.add_empty(doc)
    g11.add_normal(doc, "Версия: 1.0 | 2026-06 | Проект: greeting-service-infra")
    g11.add_empty(doc)
    add_toc(doc)

    # 14.1
    g11.add_heading2(doc, "14.1. Цель", "sec14_goal")
    g11.add_normal(
        doc,
        "1. Управлять жизненным циклом greeting-service в Kubernetes через Helm: первичный деплой, "
        "обновление образа, откат, просмотр истории release.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "2. После раздела вы умеете проверять chart (lint, template), выполнять helm upgrade --install "
        "для dev и prod, смотреть helm status/history и откатывать release при ошибке.",
    )

    # 14.2
    g11.add_heading2(doc, "14.2. Предварительные условия", "sec14_prereq")
    g11.add_normal(doc, "1. Выполнен Раздел 12: кластер доступен, Secret и registry-credentials созданы в namespace dev.")
    g11.add_empty(doc)
    g11.add_normal(doc, "2. Установлен Helm 3 на локальном ПК (проверка: helm version).")
    g11.add_empty(doc)
    g11.add_normal(doc, f"3. kubeconfig: C:\\Users\\sky\\.kube\\timeweb-greeting.yaml")
    g11.add_empty(doc)
    g11.add_normal(doc, "4. Docker-образ в Registry devtools (пример IP из terraform output):")
    g11.add_normal(doc, f"   {DEVTOOLS_IP}:5000/greeting-service:<tag>", consolas=True)
    g11.add_empty(doc)
    g11.add_normal(doc, "5. Терминал: Git Bash (Windows). kubectl и helm — клиенты к кластеру.")

    # 14.3
    g11.add_heading2(doc, "14.3. Теория: Helm и chart", "sec14_theory")
    g11.add_normal(
        doc,
        "Helm — пакетный менеджер для Kubernetes. Chart — каталог с шаблонами манифестов и файлом values.yaml. "
        "Release — установленный экземпляр chart в конкретном namespace с номером revision.",
    )
    add_external_doc_citation(
        doc,
        HELM_CHARTS_URL,
        "A chart is a collection of files that describe a related set of Kubernetes resources. "
        "A single chart might be used to deploy something simple, like a memcached pod, "
        "or something complex, like a full web app stack with HTTP servers, databases, caches, and so on.",
        "Chart — набор файлов, описывающих связанный набор ресурсов Kubernetes. "
        "Один chart может разворачивать как простой pod, так и полный стек веб-приложения.",
    )
    g11.add_empty(doc)
    g11.add_normal(doc, "Основные понятия:")
    g11.add_normal(doc, "• Chart — исходники в infra/helm/greeting-service;", consolas=True)
    g11.add_normal(doc, "• Release — greeting-service в namespace dev или prod;", consolas=True)
    g11.add_normal(doc, "• Revision — номер версии release в helm history;", consolas=True)
    g11.add_normal(doc, "• values.yaml + values-dev.yaml — параметры без правки templates/.", consolas=True)
    add_external_doc_citation(
        doc,
        INCIDENTA_HELM_URL,
        "Helm — это пакетный менеджер для Kubernetes (по смыслу как apt-get/yum), "
        "который упрощает установку, обновление и управление приложениями в кластере.",
        "Helm — пакетный менеджер для Kubernetes, упрощающий установку, обновление "
        "и управление приложениями в кластере.",
    )

    # 14.4
    g11.add_heading2(doc, "14.4. Схема работы Helm в экосистеме Kubernetes", "sec14_diagram")
    add_figure(
        doc,
        IMAGE_ECOSYSTEM,
        "Рисунок 1. Helm как компонент в экосистеме Kubernetes (источник: Incidenta)",
    )
    g11.add_normal(doc, f"Источник рисунка: {INCIDENTA_IMAGE_URL}")
    g11.add_empty(doc)
    g11.add_normal(doc, "Пояснение к рисунку 1:", consolas=True)
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "1. Helm работает поверх Kubernetes API: вы запускаете команды helm на локальном ПК или в CI/CD, "
        "а Helm Client обращается к API Server кластера и создаёт или обновляет объекты (Deployment, Service, Ingress и др.).",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "2. Chart — пакет с шаблонами и values.yaml. Release — установленный экземпляр chart в namespace. "
        "Helm хранит историю revision и позволяет откатываться без ручного kubectl apply десятков файлов.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "3. В отличие от прямого kubectl apply, Helm управляет набором ресурсов как единым релизом: "
        "upgrade, rollback и uninstall затрагивают все связанные манифесты согласованно.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "4. В нашем проекте chart greeting-service разворачивает приложение в namespace dev/prod; "
        "Secret с паролем БД создаётся отдельно (Раздел 12) и подключается через envFrom в Deployment.",
    )
    g11.add_empty(doc)
    add_figure(
        doc,
        IMAGE_PROJECT,
        "Рисунок 2. Цепочка деплоя greeting-service: chart → release → объекты в namespace dev",
    )
    g11.add_normal(doc, "Пояснение к рисунку 2:", consolas=True)
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "1. helm lint и helm template — проверка на локальном ПК без изменений в кластере.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "2. helm upgrade --install greeting-service создаёт или обновляет release; флаг --atomic откатывает "
        "release при ошибке деплоя.",
    )

    # 14.5
    g11.add_heading2(doc, "14.5. Структура chart greeting-service", "sec14_structure")
    g11.add_normal(doc, "Каталог infra/helm/greeting-service:")
    g11.add_platform_block(
        doc,
        "Структура chart",
        "Chart.yaml              # имя и версия chart\n"
        "values.yaml             # значения по умолчанию\n"
        "values-dev.yaml         # переопределения dev\n"
        "values-prod.yaml        # переопределения prod\n"
        "templates/\n"
        "  deployment.yaml       # Pod, probes, envFrom Secret\n"
        "  service.yaml          # ClusterIP :80 → :8080\n"
        "  ingress.yaml          # host, class nginx\n"
        "  serviceaccount.yaml\n"
        "  _helpers.tpl",
        yaml_block=True,
    )
    g11.add_normal(
        doc,
        f"В dev (values-dev.yaml): host={HOST_DEV}, repository={DEVTOOLS_IP}:5000/greeting-service. "
        "Тег образа — только существующий в registry (см. п. 14.7).",
    )

    # 14.6
    g11.add_heading2(doc, "14.6. Проверка chart: lint и template", "sec14_lint")
    g11.add_normal(doc, "1. Проверить синтаксис chart:")
    add_gitbash_block(doc, f"helm lint {CHART}")
    g11.add_normal(doc, "Ожидается: 1 chart(s) linted, 0 chart(s) failed.")
    g11.add_empty(doc)
    g11.add_normal(doc, "2. Посмотреть сгенерированные манифесты без apply:")
    add_gitbash_block(
        doc,
        f"helm template {RELEASE} {CHART} \\\n"
        f"  -f {CHART}/values.yaml \\\n"
        f"  -f {CHART}/values-dev.yaml \\\n"
        f"  --set image.repository={DEVTOOLS_IP}:5000/greeting-service \\\n"
        f'  --set image.tag="${{IMAGE_TAG}}"',
    )
    g11.add_normal(
        doc,
        "В выводе должны быть kind: Deployment, Service, Ingress с именем greeting-service "
        f"и host {HOST_DEV}.",
    )

    # 14.7
    g11.add_heading2(doc, "14.7. Проверка тега в registry и деплой в dev", "sec14_dev")
    g11.add_normal(
        doc,
        "1. Перед helm upgrade проверьте, какие теги реально есть в Docker Registry. "
        "Без авторизации registry вернёт UNAUTHORIZED, а не список тегов.",
    )
    add_gitbash_block(
        doc,
        f'curl -s -u docker:docker "http://{DEVTOOLS_IP}:5000/v2/greeting-service/tags/list"\n'
        f"# Пример ответа: {{\"name\":\"greeting-service\",\"tags\":[\"manual-v8\",\"manual-v7\",...]}}\n"
        f"# Тега dev-latest в registry нет — он был ошибочно указан в старой версии инструкции.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "2. Узнайте тег текущего работающего release (если деплой уже был):",
    )
    add_gitbash_block(
        doc,
        f"helm get values {RELEASE} -n {NS_DEV} | grep tag:",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "3. Задайте переменную IMAGE_TAG — любой тег из п. 1, который хотите задеплоить "
        "(например manual-v8). Не придумывайте новый тег без docker push.",
    )
    add_gitbash_block(
        doc,
        'export IMAGE_TAG=manual-v8\n'
        'echo "IMAGE_TAG=$IMAGE_TAG"',
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "4. Команда helm upgrade --install идемпотентна: первый запуск — install, следующие — upgrade.",
    )
    g11.add_empty(doc)
    g11.add_normal(doc, "5. Сначала dry-run (проверка без изменений в кластере):")
    add_gitbash_block(
        doc,
        f"helm upgrade --install {RELEASE} {CHART} \\\n"
        f"  --namespace {NS_DEV} \\\n"
        f"  --create-namespace \\\n"
        f"  -f {CHART}/values.yaml \\\n"
        f"  -f {CHART}/values-dev.yaml \\\n"
        f"  --set image.repository={DEVTOOLS_IP}:5000/greeting-service \\\n"
        f'  --set image.tag="${{IMAGE_TAG}}" \\\n'
        f"  --rollback-on-failure \\\n"
        f"  --timeout 5m \\\n"
        f"  --dry-run=client",
    )
    g11.add_empty(doc)
    g11.add_normal(doc, "6. Реальный деплой (уберите --dry-run):")
    add_gitbash_block(
        doc,
        f"helm upgrade --install {RELEASE} {CHART} \\\n"
        f"  --namespace {NS_DEV} \\\n"
        f"  --create-namespace \\\n"
        f"  -f {CHART}/values.yaml \\\n"
        f"  -f {CHART}/values-dev.yaml \\\n"
        f"  --set image.repository={DEVTOOLS_IP}:5000/greeting-service \\\n"
        f'  --set image.tag="${{IMAGE_TAG}}" \\\n'
        f"  --rollback-on-failure \\\n"
        f"  --timeout 5m",
    )
    g11.add_normal(
        doc,
        "Откуда брался несуществующий dev-latest: поле tag в values-dev.yaml (старая версия) "
        "и --set image.tag=dev-latest в инструкции. CI (ci/.gitlab-ci.yml) пушит тег "
        "CI_COMMIT_SHORT_SHA и latest, не dev-latest.",
    )
    add_external_doc_citation(
        doc,
        HELM_UPGRADE_URL,
        "This command upgrades a release to a new version of a chart.",
        "Команда обновляет release до новой версии chart.",
    )
    g11.add_normal(
        doc,
        "Флаг --rollback-on-failure: при ошибке деплоя Helm откатывает release к предыдущей revision "
        "(в Helm 4 вместо устаревшего --atomic).",
    )

    # 14.8
    g11.add_heading2(doc, "14.8. Деплой в prod", "sec14_prod")
    g11.add_normal(doc, "Аналогично dev, но namespace prod и values-prod.yaml:")
    add_gitbash_block(
        doc,
        f"helm upgrade --install {RELEASE} {CHART} \\\n"
        f"  --namespace {NS_PROD} \\\n"
        f"  --create-namespace \\\n"
        f"  -f {CHART}/values.yaml \\\n"
        f"  -f {CHART}/values-prod.yaml \\\n"
        f"  --set image.repository={DEVTOOLS_IP}:5000/greeting-service \\\n"
        f"  --set image.tag=<PROD_TAG> \\\n"
        f"  --rollback-on-failure \\\n"
        f"  --timeout 10m",
    )
    g11.add_normal(doc, "<PROD_TAG> — тег образа из pipeline или ручной сборки (не dev-latest).")

    # 14.9
    g11.add_heading2(doc, "14.9. Управление release", "sec14_manage")
    add_gitbash_block(
        doc,
        f"helm list -n {NS_DEV}\n"
        f"helm status {RELEASE} -n {NS_DEV}\n"
        f"helm history {RELEASE} -n {NS_DEV}",
    )
    g11.add_normal(doc, "Ожидается: STATUS deployed, в history — список revision с датами.")

    # 14.10
    g11.add_heading2(doc, "14.10. Откат и удаление", "sec14_rollback")
    g11.add_normal(doc, "Откат к предыдущей revision:")
    add_gitbash_block(doc, f"helm rollback {RELEASE} -n {NS_DEV}")
    g11.add_empty(doc)
    g11.add_normal(doc, "Откат к конкретной revision (номер из helm history):")
    add_gitbash_block(doc, f"helm rollback {RELEASE} 1 -n {NS_DEV}")
    g11.add_empty(doc)
    g11.add_normal(doc, "Полное удаление release и объектов chart из namespace:")
    add_gitbash_block(doc, f"helm uninstall {RELEASE} -n {NS_DEV}")
    g11.add_normal(doc, "После uninstall pod, service и ingress greeting-service в dev будут удалены.")

    # 14.11
    g11.add_heading2(doc, "14.11. Как проверить результат", "sec14_verify")
    add_gitbash_block(
        doc,
        f"helm list -n {NS_DEV}\n"
        f"helm status {RELEASE} -n {NS_DEV}\n"
        f"kubectl get pods,deployment,ingress -n {NS_DEV}\n"
        f'curl -s "http://{HOST_DEV}/api/greeting"',
    )
    g11.add_normal(doc, "Успех: helm STATUS deployed; pod Running; curl возвращает JSON с message.")

    # 14.12
    g11.add_heading2(doc, "14.12. Типичные ошибки", "sec14_errors")
    g11.add_normal(doc, "Ошибка: INSTALLATION FAILED: timed out waiting for the condition.", consolas=True)
    g11.add_normal(doc, "Причина: pod не стал Ready за --timeout (Flyway, OOM, ImagePullBackOff).", consolas=True)
    g11.add_normal(
        doc,
        "Исправление: kubectl get pods -n dev; kubectl describe pod -l app.kubernetes.io/name=greeting-service -n dev; "
        "kubectl logs -l app.kubernetes.io/name=greeting-service -n dev --tail=80. С --atomic Helm уже откатил release.",
    )
    g11.add_empty(doc)
    g11.add_normal(doc, "Ошибка: ImagePullBackOff / ErrImagePull.", consolas=True)
    g11.add_normal(doc, "Причина: --set image.tag указывает на тег, которого нет в registry.", consolas=True)
    g11.add_normal(
        doc,
        f'Исправление: curl -s -u docker:docker "http://{DEVTOOLS_IP}:5000/v2/greeting-service/tags/list" '
        "— возьмите тег из ответа; helm get values greeting-service -n dev | grep tag: — текущий тег.",
    )
    g11.add_empty(doc)
    g11.add_normal(doc, "Ошибка: Error from server (NotFound) при kubectl.", consolas=True)
    g11.add_normal(doc, "Причина: неверное имя ресурса или namespace.", consolas=True)
    g11.add_normal(doc, "Исправление: kubectl get pods,deployment,ingress -n dev", consolas=True)

    doc.save(BUILD)
    try:
        if OUTPUT.exists():
            OUTPUT.unlink()
        BUILD.replace(OUTPUT)
        saved = OUTPUT
        if BUILD.exists():
            BUILD.unlink()
    except PermissionError:
        saved = BUILD

    print(f"Written: {saved}")


if __name__ == "__main__":
    main()
