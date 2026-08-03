# -*- coding: utf-8 -*-
"""Generate docs/Razdel-13-kubernetes-manual.docx — Section 13 kubectl reference."""
from __future__ import annotations

import importlib.util
import re
import shutil
from pathlib import Path

from docx import Document

ROOT = Path(r"D:/Project_infra/greeting-service-infra")
TEMPLATE_CANDIDATES = [
    ROOT / "docs/Razdel-flyway-migrations.docx",
    ROOT / "docs/Traffic-Visibility-Caretta.docx",
    ROOT / "docs/Razdel-14-helm.docx",
]
OUTPUT = ROOT / "docs/Razdel-13-kubernetes-manual.docx"
BUILD = ROOT / "docs/_build_Razdel-13-kubernetes-manual.docx"

G11 = Path(r"C:/Users/sky/AppData/Local/Temp/gen_razdel11a_docx.py")
if not G11.exists():
    raise FileNotFoundError(f"Style module not found: {G11}")
spec = importlib.util.spec_from_file_location("g11", G11)
g11 = importlib.util.module_from_spec(spec)
assert spec.loader
spec.loader.exec_module(g11)

NS = "dev"
DEPLOY = "greeting-service"
SVC = "greeting-service"
INGRESS = "greeting-service"
SECRET = "greeting-service-secret"
LABEL = "app.kubernetes.io/name=greeting-service"
KUBECONFIG_GB = "export KUBECONFIG=/c/Users/sky/.kube/timeweb-greeting.yaml"
KUBECONFIG_WIN = r"C:\Users\sky\.kube\timeweb-greeting.yaml"
HOST_DEV = "greeting-dev.cloud-terra.online"

KUBECTL_URL = "https://kubernetes.io/docs/reference/kubectl/"
K8S_NS_URL = "https://kubernetes.io/docs/concepts/overview/working-with-objects/namespaces/"
K8S_DESCRIBE_URL = "https://kubernetes.io/docs/reference/kubectl/generated/kubectl_describe/"
K8S_LOGS_URL = "https://kubernetes.io/docs/reference/kubectl/generated/kubectl_logs/"
K8S_EXEC_URL = "https://kubernetes.io/docs/reference/kubectl/generated/kubectl_exec/"
K8S_PF_URL = "https://kubernetes.io/docs/tasks/access-application-cluster/port-forward-access-application-cluster/"
K8S_SECRET_URL = "https://kubernetes.io/docs/concepts/configuration/secret/"
K8S_DEPLOY_URL = "https://kubernetes.io/docs/concepts/workloads/controllers/deployment/"
K8S_METRICS_URL = "https://kubernetes.io/docs/tasks/debug/debug-application/resource-metrics-pipeline/"
K8S_INGRESS_URL = "https://kubernetes.io/docs/concepts/services-networking/ingress/"

TOC = [
    ("13.1. Цель", "sec13_goal"),
    ("13.2. Предварительные условия", "sec13_prereq"),
    ("13.3. Настройка kubectl на локальном ПК", "sec13_kubectl_setup"),
    ("13.4. Что делается на локальном ПК", "sec13_local"),
    ("13.5. Навигация по кластеру", "sec13_nav"),
    ("13.6. Детальная информация и события", "sec13_describe"),
    ("13.7. Работа с логами", "sec13_logs"),
    ("13.8. Выполнение команд в поде", "sec13_exec"),
    ("13.9. Port-forward без Ingress", "sec13_pf"),
    ("13.10. Secrets и ConfigMaps", "sec13_secrets"),
    ("13.11. Управление Deployment", "sec13_deploy"),
    ("13.12. Просмотр ресурсов узлов", "sec13_top"),
    ("13.13. Работа с Ingress", "sec13_ingress"),
    ("13.14. Что делается на сервере", "sec13_server"),
    ("13.15. Как проверить результат", "sec13_verify"),
    ("13.16. Типичные ошибки", "sec13_errors"),
]


def add_external_doc_citation(doc: Document, url: str, quote_en: str, quote_ru: str) -> None:
    g11.add_empty(doc)
    g11.add_normal(doc, f"Источник: {url}")
    g11.add_empty(doc)
    g11.add_normal(doc, "Цитата:")
    g11.add_platform_block(doc, "Оригинал (English)", quote_en)
    g11.add_empty(doc)
    g11.add_normal(doc, "Перевод:")
    g11.add_platform_block(doc, "Перевод на русский", quote_ru)


def add_toc(doc: Document) -> None:
    g11.add_heading2(doc, "Оглавление", "sec13_toc")
    for title, anchor in TOC:
        indent = bool(re.match(r"13\.\d+\.\d+", title))
        g11.add_hyperlink_paragraph(doc, title, anchor, indent=indent)
    g11.add_empty(doc)


def resolve_template() -> Path:
    for path in TEMPLATE_CANDIDATES:
        if path.exists():
            return path
    raise FileNotFoundError("No Word template found in docs/")


def add_gitbash_block(doc: Document, code: str) -> None:
    g11.add_platform_block(
        doc,
        "Локальный ПК — Windows (Git Bash), корень репозитория",
        f"cd '/d/Project_infra/greeting-service-infra'\n{KUBECONFIG_GB}\n{code}",
    )


def main() -> None:
    g11._bookmark_id = 0
    template = resolve_template()
    shutil.copy2(template, BUILD)
    doc = Document(BUILD)
    g11.clear_document_body(doc)

    g11.add_normal(doc, "Раздел 13. Ручное управление Kubernetes с локального ПК", consolas=True)
    g11.add_empty(doc)
    g11.add_normal(doc, "Версия: 1.0 | 2026-06 | Проект: greeting-service-infra")
    g11.add_empty(doc)
    add_toc(doc)

    # --- 13.1 ---
    g11.add_heading2(doc, "13.1. Цель", "sec13_goal")
    g11.add_normal(
        doc,
        "1. Собрать в одном разделе практический справочник команд kubectl для ежедневной работы "
        "с Kubernetes-кластером greeting-service из локальной среды разработчика.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "2. После изучения и выполнения проверок из п. 13.15 вы сможете: просматривать состояние "
        "кластера и namespace, читать логи и события, проверять Secret, выполнять rolling restart, "
        "диагностировать Ingress и временно открывать приложение через port-forward — без входа "
        "на worker-узлы по SSH.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "3. Раздел 13 не заменяет первичный деплой (Раздел 12) и не описывает Helm (Раздел 14). "
        "Здесь — только повседневное управление уже развёрнутым приложением.",
    )
    add_external_doc_citation(
        doc,
        KUBECTL_URL,
        "kubectl is a command-line tool for controlling Kubernetes clusters. "
        "It is the primary way to interact with a cluster from your local machine.",
        "kubectl — инструмент командной строки для управления кластерами Kubernetes. "
        "Это основной способ взаимодействия с кластером с локального компьютера.",
    )

    # --- 13.2 ---
    g11.add_heading2(doc, "13.2. Предварительные условия", "sec13_prereq")
    g11.add_normal(doc, "1. Выполнен Раздел 9: terraform apply завершён, managed Kubernetes-кластер создан.")
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        f"2. Выполнен Раздел 12: приложение greeting-service развёрнуто в namespace {NS}, "
        "Helm release установлен, pod в статусе Running.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        f"3. Файл kubeconfig сохранён: {KUBECONFIG_WIN} "
        "(скрипт scripts/get-kubeconfig.sh после terraform apply).",
    )
    g11.add_empty(doc)
    g11.add_normal(doc, "4. На локальном ПК установлен kubectl (клиент) — версия должна быть совместима с кластером.")
    g11.add_empty(doc)
    g11.add_normal(doc, "5. Роли терминалов в этом разделе:")
    g11.add_normal(doc, "- kubectl, curl, base64 — Git Bash (Windows);", consolas=True)
    g11.add_normal(
        doc,
        f"- kubectl get nodes с явным --kubeconfig {KUBECONFIG_WIN} — допустимо из Windows cmd;",
        consolas=True,
    )
    g11.add_normal(doc, "- terraform output — только WSL Ubuntu (не для команд kubectl).", consolas=True)
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        f"6. В примерах команд используется namespace {NS} (dev-окружение). "
        "Для prod замените -n dev на -n prod и host в curl на greeting.cloud-terra.online.",
    )

    # --- 13.3 ---
    g11.add_heading2(doc, "13.3. Настройка kubectl на локальном ПК", "sec13_kubectl_setup")
    g11.add_normal(
        doc,
        "1. kubectl — клиент: он отправляет запросы в API-сервер кластера. "
        "Адрес API и учётные данные хранятся в файле kubeconfig.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "2. Перед любой командой kubectl в Git Bash задайте переменную KUBECONFIG. "
        "Без неё kubectl может обращаться к localhost:8080 и вернуть ошибку connection refused.",
    )
    add_gitbash_block(
        doc,
        "kubectl cluster-info\nkubectl get nodes",
    )
    g11.add_normal(
        doc,
        "3. Успех: cluster-info выводит URL API-сервера Timeweb Cloud; get nodes показывает worker-узлы "
        "в статусе Ready.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "4. Чтобы не вводить export каждый раз, добавьте строку KUBECONFIG в ~/.bashrc в Git Bash "
        "(см. Раздел 12, п. 12.10.1).",
    )
    add_external_doc_citation(
        doc,
        KUBECTL_URL,
        "By default, kubectl looks for a file named config in the $HOME/.kube directory. "
        "You can specify other kubeconfig files by setting the KUBECONFIG environment variable "
        "or by setting the --kubeconfig flag.",
        "По умолчанию kubectl ищет файл config в каталоге $HOME/.kube. "
        "Другой kubeconfig можно указать переменной окружения KUBECONFIG или флагом --kubeconfig.",
    )

    # --- 13.4 ---
    g11.add_heading2(doc, "13.4. Что делается на локальном ПК", "sec13_local")
    g11.add_normal(
        doc,
        "1. Все команды раздела 13 выполняются на локальном ПК разработчика с корректно "
        "настроенным KUBECONFIG. kubectl не требует SSH на devtools или worker-узлы.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "2. Типовой порядок диагностики проблемы с приложением:",
    )
    g11.add_normal(doc, "• kubectl get pods -n dev — pod Running или нет;", consolas=True)
    g11.add_normal(doc, "• kubectl describe pod … — раздел Events внизу;", consolas=True)
    g11.add_normal(doc, "• kubectl logs … — вывод Spring Boot / Flyway;", consolas=True)
    g11.add_normal(doc, "• kubectl get ingress -n dev — host и адрес;", consolas=True)
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "3. Имена ресурсов в кластере после Helm release greeting-service:",
    )
    g11.add_normal(doc, f"• Deployment: {DEPLOY}", consolas=True)
    g11.add_normal(doc, f"• Service: {SVC}", consolas=True)
    g11.add_normal(doc, f"• Secret приложения: {SECRET}", consolas=True)
    g11.add_normal(doc, f"• Label selector: {LABEL}", consolas=True)

    # --- 13.5 ---
    g11.add_heading2(doc, "13.5. Навигация по кластеру", "sec13_nav")
    g11.add_normal(
        doc,
        "1. Команды get выводят краткий список ресурсов. Это первый шаг при любой проверке.",
    )
    add_gitbash_block(
        doc,
        "kubectl cluster-info\n"
        "kubectl get nodes\n"
        "kubectl get nodes -o wide\n"
        "kubectl get all --all-namespaces\n"
        f"kubectl get all -n {NS}\n"
        f"kubectl get pods -n {NS}\n"
        f"kubectl get pods -n {NS} -o wide\n"
        f"# Наблюдение в реальном времени (Ctrl+C для выхода):\n"
        f"kubectl get pods -n {NS} -w",
    )
    g11.add_normal(doc, "2. Расшифровка ключевых команд:")
    g11.add_normal(
        doc,
        "• cluster-info — адрес API-сервера и адрес CoreDNS; подтверждает, что клиент "
        "подключён к нужному кластеру.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "• get nodes -o wide — дополнительно показывает INTERNAL-IP, EXTERNAL-IP и версию kubelet. "
        "EXTERNAL-IP worker-узла в Timeweb Cloud используется как INGRESS_IP (Раздел 12).",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        f"• get all -n {NS} — сводка Deployment, ReplicaSet, Pod, Service, Ingress в одном выводе.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        f"• get pods -n {NS} -w — режим watch: таблица обновляется при изменении статуса pod "
        "(полезно после helm upgrade или rollout restart).",
    )
    add_external_doc_citation(
        doc,
        K8S_NS_URL,
        "Namespaces provide a scope for names. Names of resources need to be unique within a namespace, "
        "but not across namespaces.",
        "Namespace задаёт область имён. Имена ресурсов должны быть уникальны внутри namespace, "
        "но не обязаны быть уникальны между разными namespace.",
    )

    # --- 13.6 ---
    g11.add_heading2(doc, "13.6. Детальная информация и события", "sec13_describe")
    g11.add_normal(
        doc,
        "1. kubectl describe выводит подробное описание ресурса и связанных событий. "
        "При диагностике CrashLoopBackOff, ImagePullBackOff, OOMKilled — это основная команда.",
    )
    add_gitbash_block(
        doc,
        f"POD=$(kubectl get pods -n {NS} -l {LABEL} -o jsonpath='{{.items[0].metadata.name}}')\n"
        f"echo \"Pod: $POD\"\n"
        f"kubectl describe pod \"$POD\" -n {NS}\n"
        f"kubectl describe deployment {DEPLOY} -n {NS}\n"
        f"kubectl describe ingress {INGRESS} -n {NS}",
    )
    g11.add_normal(doc, "2. На что смотреть в выводе describe pod:")
    g11.add_normal(doc, "• Status / Reason — OOMKilled, Error, Completed;", consolas=True)
    g11.add_normal(doc, "• State: Waiting — причина (CrashLoopBackOff, ImagePullBackOff);", consolas=True)
    g11.add_normal(doc, "• Limits / Requests — memory и CPU (OOM при превышении limit);", consolas=True)
    g11.add_normal(doc, "• Events — хронология: Scheduled, Pulled, Created, Started, Killing.", consolas=True)
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "3. describe deployment показывает стратегию RollingUpdate, число replicas и условия Available.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "4. describe ingress показывает правило host → service:port и адрес, на который "
        "направляется внешний трафик.",
    )
    add_external_doc_citation(
        doc,
        K8S_DESCRIBE_URL,
        "Print a detailed description of the selected resources, including related resources "
        "such as events or controllers.",
        "Вывести подробное описание выбранных ресурсов, включая связанные ресурсы, "
        "такие как events (события) или controllers (контроллеры).",
    )

    # --- 13.7 ---
    g11.add_heading2(doc, "13.7. Работа с логами", "sec13_logs")
    g11.add_normal(
        doc,
        "1. Логи приложения (stdout/stderr контейнера) читаются командой kubectl logs. "
        "Для Spring Boot здесь видны Flyway, Tomcat, ошибки JDBC.",
    )
    add_gitbash_block(
        doc,
        f"POD=$(kubectl get pods -n {NS} -l {LABEL} -o jsonpath='{{.items[0].metadata.name}}')\n"
        f"# Последние 100 строк:\n"
        f"kubectl logs \"$POD\" -n {NS} --tail=100\n"
        f"# Режим слежения (Ctrl+C для выхода):\n"
        f"kubectl logs -f \"$POD\" -n {NS}\n"
        f"# Логи предыдущего (упавшего) контейнера:\n"
        f"kubectl logs \"$POD\" -n {NS} --previous\n"
        f"# По label — все pod deployment:\n"
        f"kubectl logs -l {LABEL} -n {NS} --tail=50",
    )
    g11.add_normal(doc, "2. Когда использовать --previous:")
    g11.add_normal(
        doc,
        "• Pod в CrashLoopBackOff: текущий контейнер может ещё не успеть записать лог; "
        "--previous показывает вывод завершившегося экземпляра.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "• OOMKilled: в --previous часто видно Started GreetingServiceApplication, "
        "после чего pod перезапускается — признак малого memory limit.",
    )
    add_external_doc_citation(
        doc,
        K8S_LOGS_URL,
        "Print the logs for a container in a pod or specified resource. "
        "If the pod has only one container, the container name is optional.",
        "Вывести логи контейнера в pod или указанного ресурса. "
        "Если в pod один контейнер, имя контейнера можно не указывать.",
    )

    # --- 13.8 ---
    g11.add_heading2(doc, "13.8. Выполнение команд в поде", "sec13_exec")
    g11.add_normal(
        doc,
        "1. kubectl exec выполняет команду внутри работающего контейнера. "
        "Используется для проверки переменных окружения и health endpoint без Ingress.",
    )
    add_gitbash_block(
        doc,
        f"POD=$(kubectl get pods -n {NS} -l {LABEL} -o jsonpath='{{.items[0].metadata.name}}')\n"
        f"# Интерактивная оболочка (образ Spring Boot — /bin/sh):\n"
        f"kubectl exec -it \"$POD\" -n {NS} -- /bin/sh\n"
        f"# Переменные окружения (проверка Secret):\n"
        f"kubectl exec \"$POD\" -n {NS} -- env | grep -E 'DB_|SPRING'\n"
        f"# Health изнутри pod:\n"
        f"kubectl exec \"$POD\" -n {NS} -- wget -qO- http://localhost:8080/actuator/health",
    )
    g11.add_normal(
        doc,
        "2. Образ greeting-service может не содержать curl; для проверки health используйте wget "
        "или exec -it с /bin/sh и встроенными утилитами.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "3. env | grep DB_ показывает, попали ли DB_URL и DB_PASSWORD из Secret greeting-service-secret. "
        "Значения не должны быть пустыми.",
    )
    add_external_doc_citation(
        doc,
        K8S_EXEC_URL,
        "Execute a command in a container.",
        "Выполнить команду в контейнере.",
    )

    # --- 13.9 ---
    g11.add_heading2(doc, "13.9. Port-forward без Ingress", "sec13_pf")
    g11.add_normal(
        doc,
        "1. Port-forward создаёт туннель с локального ПК на pod или Service в кластере. "
        "Приложение временно доступно на localhost — без DNS и Ingress.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "2. Терминал с port-forward должен оставаться открытым. Для проверки откройте второй "
        "терминал Git Bash.",
    )
    add_gitbash_block(
        doc,
        f"# Терминал 1 — оставить открытым:\n"
        f"kubectl port-forward svc/{SVC} 8080:80 -n {NS}\n"
        f"# Терминал 2 — проверка:\n"
        f"curl -s http://localhost:8080/api/greeting",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "3. Схема трафика: localhost:8080 → Service :80 → Pod :8080. "
        "Service greeting-service слушает порт 80 и направляет на containerPort 8080.",
    )
    g11.add_empty(doc)
    g11.add_normal(doc, "4. Проброс напрямую к pod (минуя Service):")
    add_gitbash_block(
        doc,
        f"POD=$(kubectl get pods -n {NS} -l {LABEL} -o jsonpath='{{.items[0].metadata.name}}')\n"
        f"kubectl port-forward pod/\"$POD\" 8080:8080 -n {NS}",
    )
    add_external_doc_citation(
        doc,
        K8S_PF_URL,
        "Connections made to local port 28015 are forwarded to port 27017 of the Pod that is running "
        "the MongoDB server. With this connection in place, you can use your local workstation to "
        "debug the database that is running in the Pod.",
        "Соединения на локальный порт 28015 перенаправляются на порт 27017 pod, в котором "
        "работает сервер MongoDB. С таким соединением можно с локального компьютера отлаживать "
        "приложение (или сервис), работающее в pod.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "5. После запуска port-forward сервис временно доступен по адресу "
        "http://localhost:8080/api/greeting — это ожидаемый результат проверки.",
    )

    # --- 13.10 ---
    g11.add_heading2(doc, "13.10. Secrets и ConfigMaps", "sec13_secrets")
    g11.add_normal(
        doc,
        "1. Secret greeting-service-secret хранит DB_URL и DB_PASSWORD. "
        "Deployment подключает его через envFrom (см. Helm chart).",
    )
    add_gitbash_block(
        doc,
        f"kubectl get secrets -n {NS}\n"
        f"kubectl get secret {SECRET} -n {NS} -o yaml\n"
        f"# Декодировать DB_URL:\n"
        f"kubectl get secret {SECRET} -n {NS} -o jsonpath='{{.data.DB_URL}}' | base64 -d\n"
        f"echo \"\"\n"
        f"kubectl get configmap -n {NS}",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "2. Значения в yaml выводятся в base64 — это не шифрование, а кодировка. "
        "base64 -d показывает фактическую строку JDBC.",
    )
    g11.add_empty(doc)
    g11.add_normal(doc, "3. Обновить Secret без ошибки already exists (идемпотентно):")
    add_gitbash_block(
        doc,
        f"kubectl create secret generic {SECRET} \\\n"
        f"  --namespace {NS} \\\n"
        f'  --from-literal=DB_URL="jdbc:postgresql://10.10.0.5:5432/greeting_db" \\\n'
        f'  --from-literal=DB_PASSWORD="ВАШ_ПАРОЛЬ" \\\n'
        f"  --dry-run=client -o yaml | kubectl apply -f -",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "4. После изменения Secret выполните rollout restart deployment — pod перечитает переменные "
        "(см. п. 13.11).",
    )
    add_external_doc_citation(
        doc,
        K8S_SECRET_URL,
        "Kubernetes Secrets let you store and manage sensitive information, such as passwords, "
        "OAuth tokens, and ssh keys.",
        "Kubernetes Secrets позволяют хранить и управлять чувствительной информацией: "
        "паролями, OAuth-токенами, SSH-ключами.",
    )

    # --- 13.11 ---
    g11.add_heading2(doc, "13.11. Управление Deployment", "sec13_deploy")
    g11.add_normal(
        doc,
        f"1. Deployment управляет ReplicaSet и pod. Helm release greeting-service создаёт Deployment {DEPLOY}.",
    )
    add_gitbash_block(
        doc,
        f"# Масштабирование (dev обычно 1 replica):\n"
        f"kubectl scale deployment {DEPLOY} -n {NS} --replicas=2\n"
        f"# Rolling restart — пересоздать pod без helm:\n"
        f"kubectl rollout restart deployment/{DEPLOY} -n {NS}\n"
        f"kubectl rollout status deployment/{DEPLOY} -n {NS}\n"
        f"kubectl rollout history deployment/{DEPLOY} -n {NS}\n"
        f"# Откат на предыдущую revision:\n"
        f"kubectl rollout undo deployment/{DEPLOY} -n {NS}",
    )
    g11.add_normal(doc, "2. Когда нужен rollout restart:")
    g11.add_normal(doc, "• после изменения Secret или ConfigMap;", consolas=True)
    g11.add_normal(doc, "• после смены memory limits в values-dev.yaml и helm upgrade;", consolas=True)
    g11.add_normal(doc, "• для принудительного pull образа с tag dev-latest.", consolas=True)
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "3. rollout restart не меняет образ — только пересоздаёт pod. "
        "Для смены image tag используйте helm upgrade (Раздел 14).",
    )
    add_external_doc_citation(
        doc,
        K8S_DEPLOY_URL,
        "A Deployment provides declarative updates for Pods and ReplicaSets. "
        "You describe a desired state in a Deployment, and the Deployment Controller "
        "changes the actual state to the desired state at a controlled rate.",
        "Deployment обеспечивает декларативные обновления Pod и ReplicaSet. "
        "Вы описываете желаемое состояние в Deployment, а контроллер Deployment "
        "постепенно приводит фактическое состояние к желаемому.",
    )

    # --- 13.12 ---
    g11.add_heading2(doc, "13.12. Просмотр ресурсов узлов", "sec13_top")
    g11.add_normal(
        doc,
        "1. kubectl top показывает текущее потребление CPU и памяти. "
        "Требует работающий metrics-server в кластере (в Timeweb Cloud обычно предустановлен).",
    )
    add_gitbash_block(
        doc,
        "kubectl top nodes\n"
        f"kubectl top pods -n {NS}",
    )
    g11.add_normal(
        doc,
        "2. Если top pods показывает память pod близко к limit — риск OOMKilled. "
        "Для Spring Boot в dev рабочий minimum limit — 512Mi (values-dev.yaml).",
    )
    add_external_doc_citation(
        doc,
        K8S_METRICS_URL,
        "Resource usage metrics, such as container CPU and memory usage, are available in Kubernetes "
        "through the Metrics API. These metrics can be accessed either directly by user, "
        "by using kubectl top command, or indirectly by the cluster's Horizontal Pod Autoscaler.",
        "Метрики использования ресурсов (CPU и память контейнеров) доступны в Kubernetes "
        "через Metrics API. К ним можно обратиться напрямую, командой kubectl top, "
        "или косвенно — через Horizontal Pod Autoscaler кластера.",
    )

    # --- 13.13 ---
    g11.add_heading2(doc, "13.13. Работа с Ingress", "sec13_ingress")
    g11.add_normal(
        doc,
        "1. Ingress маршрутизирует HTTP-запросы с внешнего host на Service приложения.",
    )
    add_gitbash_block(
        doc,
        f"kubectl get ingress --all-namespaces\n"
        f"kubectl get ingress -n {NS}\n"
        f"kubectl describe ingress {INGRESS} -n {NS}\n"
        f"# Логи NGINX Ingress Controller (если namespace ingress-nginx есть):\n"
        f"kubectl logs -n ingress-nginx -l app.kubernetes.io/name=ingress-nginx --tail=100",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        f"2. В dev host Ingress: {HOST_DEV}. "
        "Проверка снаружи через DNS (после настройки A-записи, Раздел 12):",
    )
    add_gitbash_block(
        doc,
        f'curl -s "http://{HOST_DEV}/api/greeting"',
    )
    g11.add_normal(
        doc,
        "Ожидается: JSON с полем message (HTTP 200). Команда проверена в Git Bash.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "3. Важно для Timeweb Cloud: kubectl get svc -n ingress-nginx часто возвращает пустой список "
        "или не показывает EXTERNAL-IP — Ingress Controller работает на worker через hostNetwork.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "4. Внешний IP для DNS A-записи берите из EXTERNAL-IP worker-узла (не из ingress-nginx Service):",
    )
    add_gitbash_block(
        doc,
        "kubectl get nodes -o wide\n"
        "INGRESS_IP=$(kubectl get nodes -o jsonpath='{.items[0].status.addresses[?(@.type==\"ExternalIP\")].address}')\n"
        'echo "INGRESS_IP=$INGRESS_IP"',
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "5. Альтернатива curl по DNS — запрос по IP с заголовком Host (если DNS ещё не обновился):",
    )
    add_gitbash_block(
        doc,
        f'curl -s -H "Host: {HOST_DEV}" "http://$INGRESS_IP/api/greeting"',
    )
    add_external_doc_citation(
        doc,
        K8S_INGRESS_URL,
        "An Ingress may be configured to give Services externally-reachable URLs, "
        "load balance traffic, terminate SSL / TLS, and offer name based virtual hosting.",
        "Ingress может быть настроен так, чтобы давать Service URL, доступные извне, "
        "балансировать трафик, завершать SSL/TLS и обеспечивать виртуальный хостинг по имени.",
    )

    # --- 13.14 ---
    g11.add_heading2(doc, "13.14. Что делается на сервере", "sec13_server")
    g11.add_normal(
        doc,
        "1. В рамках раздела 13 команды на devtools-сервере и worker-узлах не выполняются. "
        "Управление идёт через Kubernetes API с локального ПК.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "2. В облаке уже работают: control plane managed Kubernetes, worker-узлы, "
        "Ingress Controller (ingress-nginx), Deployment, Service, Secret и Ingress greeting-service.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "3. Прямой SSH на worker в Timeweb Cloud для повседневной работы не требуется. "
        "Исключение — редкая низкоуровневая диагностика (п. 12.7.1, insecure registry через kubectl debug).",
    )

    # --- 13.15 ---
    g11.add_heading2(doc, "13.15. Как проверить результат", "sec13_verify")
    g11.add_normal(
        doc,
        "1. Базовая проверка: локальный ПК видит кластер и ресурсы приложения в namespace dev.",
    )
    add_gitbash_block(
        doc,
        "kubectl cluster-info\n"
        "kubectl get nodes\n"
        f"kubectl get pods -n {NS}\n"
        f"kubectl get ingress -n {NS}",
    )
    g11.add_normal(
        doc,
        "2. Успех: команды выполняются без connection refused и без Error from server (Forbidden); "
        "pod в STATUS Running; ingress содержит host greeting-dev.cloud-terra.online.",
    )
    g11.add_empty(doc)
    g11.add_normal(doc, "3. Проверка describe и logs:")
    add_gitbash_block(
        doc,
        f"POD=$(kubectl get pods -n {NS} -l {LABEL} -o jsonpath='{{.items[0].metadata.name}}')\n"
        f"kubectl describe pod \"$POD\" -n {NS} | tail -20\n"
        f"kubectl logs \"$POD\" -n {NS} --tail=30",
    )
    g11.add_normal(
        doc,
        "4. В Events не должно быть повторяющихся Back-off или Failed; в logs — "
        "Started GreetingServiceApplication.",
    )
    g11.add_empty(doc)
    g11.add_normal(doc, "5. Проверка port-forward (терминал 1 — port-forward, терминал 2 — curl):")
    add_gitbash_block(
        doc,
        f"kubectl port-forward svc/{SVC} 8080:80 -n {NS}\n"
        "# Во втором терминале:\n"
        "curl -s http://localhost:8080/api/greeting",
    )
    g11.add_empty(doc)
    g11.add_normal(doc, "6. Проверка Ingress через DNS:")
    add_gitbash_block(
        doc,
        f'curl -s "http://{HOST_DEV}/api/greeting"',
    )
    g11.add_normal(
        doc,
        "7. Ожидаемый ответ curl — JSON с полем message. "
        "Если port-forward работает, а Ingress снаружи — нет, проверьте DNS A-запись и host в Ingress.",
    )

    # --- 13.16 ---
    g11.add_heading2(doc, "13.16. Типичные ошибки", "sec13_errors")
    g11.add_normal(doc, "Ошибка: The connection to the server was refused или no such host.", consolas=True)
    g11.add_normal(doc, "Причина: не задан KUBECONFIG, устаревший kubeconfig или нет сети до API.", consolas=True)
    g11.add_normal(
        doc,
        f"Исправление: выполните {KUBECONFIG_GB}, проверьте файл {KUBECONFIG_WIN}, "
        "при необходимости повторите scripts/get-kubeconfig.sh.",
    )
    g11.add_empty(doc)
    g11.add_normal(doc, "Ошибка: Error from server (NotFound).", consolas=True)
    g11.add_normal(doc, "Причина: неверное имя ресурса или namespace.", consolas=True)
    g11.add_normal(
        doc,
        f"Исправление: kubectl get pods,deployment,ingress -n {NS} — возьмите точное имя из вывода.",
    )
    g11.add_empty(doc)
    g11.add_normal(doc, "Ошибка: get svc -n ingress-nginx (без kubectl).", consolas=True)
    g11.add_normal(doc, "Причина: kubectl — отдельная команда, не подкоманда shell.", consolas=True)
    g11.add_normal(doc, "Исправление: kubectl get svc -n ingress-nginx", consolas=True)
    g11.add_empty(doc)
    g11.add_normal(doc, "Ошибка: kubectl top не возвращает метрики.", consolas=True)
    g11.add_normal(doc, "Причина: metrics-server недоступен или нет прав RBAC.", consolas=True)
    g11.add_normal(
        doc,
        "Исправление: kubectl get pods -n kube-system | grep metrics — pod должен быть Running.",
    )
    g11.add_empty(doc)
    g11.add_normal(doc, "Ошибка: port-forward не отвечает на localhost.", consolas=True)
    g11.add_normal(doc, "Причина: pod не Ready, неверный порт Service или port-forward закрыт.", consolas=True)
    g11.add_normal(
        doc,
        f"Исправление: kubectl get pods -n {NS}; kubectl get svc {SVC} -n {NS}; "
        "убедитесь, что port-forward запущен и терминал не закрыт.",
    )
    g11.add_empty(doc)
    g11.add_normal(doc, "Ошибка: get svc -n ingress-nginx без EXTERNAL-IP.", consolas=True)
    g11.add_normal(doc, "Причина: на Timeweb Cloud Ingress не получает отдельный LoadBalancer IP.", consolas=True)
    g11.add_normal(
        doc,
        "Исправление: используйте EXTERNAL-IP worker-узла (kubectl get nodes -o wide), "
        "не ждите EXTERNAL-IP у ingress-nginx-controller.",
    )
    g11.add_empty(doc)
    g11.add_normal(doc, "Ошибка: CrashLoopBackOff / OOMKilled.", consolas=True)
    g11.add_normal(doc, "Причина: приложение падает при старте или превышен memory limit.", consolas=True)
    g11.add_normal(
        doc,
        "Исправление: kubectl logs --previous; kubectl describe pod | grep -A5 Limits; "
        "для OOM — увеличить memory limits в values-dev.yaml и helm upgrade (Раздел 12, 14).",
    )

    doc.save(BUILD)
    try:
        if OUTPUT.exists():
            OUTPUT.unlink()
        BUILD.replace(OUTPUT)
        saved = OUTPUT
        if BUILD.exists():
            BUILD.unlink()
    except PermissionError:
        saved = BUILD

    print(f"Written: {saved}")


if __name__ == "__main__":
    main()
