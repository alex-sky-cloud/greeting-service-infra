# -*- coding: utf-8 -*-
"""Generate docs/Caretta-Radar-Grafana-ispolzovanie.docx"""
from __future__ import annotations

import importlib.util
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches

ROOT = Path(r"D:/Project_infra/greeting-service-infra")
TEMPLATE = ROOT / "docs/Traffic-Visibility-Caretta.docx"
if not TEMPLATE.exists():
    TEMPLATE = ROOT / "docs/Razdel-flyway-migrations.docx"
OUTPUT = ROOT / "docs/Caretta-Radar-Grafana-ispolzovanie.docx"
DIAGRAM_SCRIPT = ROOT / "docs/caretta_usage_diagram_render.py"
IMAGE = ROOT / "docs/images/traffic-visibility/caretta-radar-grafana.png"

G11 = Path(r"C:/Users/sky/AppData/Local/Temp/gen_razdel11a_docx.py")
spec = importlib.util.spec_from_file_location("g11", G11)
g11 = importlib.util.module_from_spec(spec)
assert spec.loader
spec.loader.exec_module(g11)

RADAR_TRAFFIC_URL = "https://radarhq.io/docs/features/traffic"
CARETTA_URL = "https://github.com/groundcover-com/caretta"

TOC = [
    ("1. Суть проблемы", "sec_cu_1"),
    ("2. Что установлено в кластере", "sec_cu_2"),
    ("3. Почему Radar Traffic не работает", "sec_cu_3"),
    ("4. Grafana: откуда она и почему не видна", "sec_cu_4"),
    ("5. Схема использования имеющихся компонентов", "sec_cu_5"),
    ("6. Пошаговая инструкция: Caretta через Grafana", "sec_cu_6"),
    ("7. Генерация трафика для карты", "sec_cu_7"),
    ("8. Что использовать Radar", "sec_cu_8"),
    ("9. Проверка состояния Caretta", "sec_cu_9"),
    ("10. Итог", "sec_cu_10"),
]


def add_external_doc_citation(doc: Document, url: str, quote_en: str, quote_ru: str) -> None:
    g11.add_empty(doc)
    g11.add_normal(doc, f"Источник: {url}")
    g11.add_empty(doc)
    g11.add_normal(doc, "Цитата:")
    g11.add_platform_block(doc, "Оригинал (English)", quote_en)
    g11.add_empty(doc)
    g11.add_normal(doc, "Перевод:")
    g11.add_platform_block(doc, "Перевод на русский", quote_ru)


def add_toc(doc: Document) -> None:
    g11.add_heading2(doc, "Оглавление", "sec_cu_toc")
    for title, anchor in TOC:
        g11.add_hyperlink_paragraph(doc, title, anchor, indent=False)
    g11.add_empty(doc)


def add_figure(doc: Document, path: Path, caption: str) -> None:
    g11.add_empty(doc)
    g11.add_normal(doc, caption, consolas=True)
    g11.add_empty(doc)
    doc.add_picture(str(path), width=Inches(6.5))
    g11.add_empty(doc)


def ensure_diagram() -> None:
    if not IMAGE.exists():
        subprocess.run([sys.executable, str(DIAGRAM_SCRIPT)], check=True)


def main() -> None:
    ensure_diagram()
    g11._bookmark_id = 0
    shutil.copy2(TEMPLATE, OUTPUT)
    doc = Document(OUTPUT)
    g11.clear_document_body(doc)

    g11.add_normal(
        doc,
        "Caretta после установки: почему Radar Traffic не работает и как пользоваться Grafana",
        consolas=True,
    )
    g11.add_empty(doc)
    g11.add_normal(doc, "Версия: 1.0 | 2026-06 | Кластер: twc-greeting-service-k8s, namespace dev")
    g11.add_empty(doc)
    add_toc(doc)

    # 1
    g11.add_heading2(doc, "1. Суть проблемы", "sec_cu_1")
    g11.add_normal(
        doc,
        "1. Вы установили Caretta через Helm (кнопка в Radar или команду helm install caretta). "
        "Pod caretta-* в статусе Running, но вкладка Traffic Visibility в Radar по-прежнему "
        "предлагает «Install caretta with Helm» и не показывает live-трафик.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "2. Это не означает, что Caretta сломан. Означает, что Radar и chart groundcover/caretta — "
        "разные способы интеграции. Caretta работает; смотреть карту трафика нужно через Grafana "
        "из того же Helm release, а не через Radar Traffic.",
    )

    # 2
    g11.add_heading2(doc, "2. Что установлено в кластере", "sec_cu_2")
    g11.add_normal(doc, "После helm install caretta (release caretta, namespace dev) создаются:")
    g11.add_empty(doc)
    g11.add_normal(doc, "• caretta — DaemonSet, eBPF-агент на каждом worker-узле;", consolas=True)
    g11.add_normal(doc, "• caretta-vm — Victoria Metrics, хранит метрики caretta_links_observed;", consolas=True)
    g11.add_normal(doc, "• caretta-grafana — Grafana с преднастроенным дашбордом Caretta;", consolas=True)
    g11.add_normal(doc, "• greeting-service — ваше приложение (отдельный Helm, тот же namespace dev).", consolas=True)
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "Grafana не ставилась отдельно — она входит в chart Caretta как зависимость (subchart grafana). "
        "Service caretta-grafana имеет тип ClusterIP: снаружи кластера URL по умолчанию нет.",
    )

    # 3
    g11.add_heading2(doc, "3. Почему Radar Traffic не работает", "sec_cu_3")
    g11.add_normal(
        doc,
        "Radar при старте ищет источник трафика в таком порядке: Hubble (Cilium) → Caretta → Istio. "
        "Если источник не найден, режим Traffic скрыт или показывается мастер установки.",
    )
    add_external_doc_citation(
        doc,
        RADAR_TRAFFIC_URL,
        "Caretta. Detected via the Caretta operator and its CRDs.",
        "Caretta. Обнаруживается через Caretta operator и его CRD.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "Chart groundcover/caretta устанавливает DaemonSet + Victoria Metrics + Grafana. "
        "Он не создаёт Caretta operator и CRD. Поэтому Radar не считает Caretta «установленным» "
        "для Traffic, даже когда pod caretta-l62jh и caretta-mcl6t уже Running.",
    )
    g11.add_empty(doc)
    g11.add_normal(doc, "Дополнительно:")
    g11.add_normal(doc, "• CNI кластера — Flannel, не Cilium → Hubble недоступен;", consolas=True)
    g11.add_normal(doc, "• Istio не установлен;", consolas=True)
    g11.add_normal(doc, "• namespace dev вместо caretta на Radar не влияет — нет CRD.", consolas=True)

    # 4
    g11.add_heading2(doc, "4. Grafana: откуда она и почему не видна", "sec_cu_4")
    g11.add_normal(
        doc,
        "1. Grafana появилась автоматически при helm install caretta. Проверка:",
    )
    g11.add_platform_block(
        doc,
        "Локальный ПК — Windows (Git Bash)",
        "export KUBECONFIG=/c/Users/sky/.kube/timeweb-greeting.yaml\n"
        "kubectl get pods -n dev | grep grafana\n"
        "kubectl get svc -n dev caretta-grafana",
    )
    g11.add_normal(
        doc,
        "2. Ожидаемый вывод: pod caretta-grafana-* Running; Service ClusterIP, порт 80/TCP.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "3. Grafana не открывается в браузере сама по себе — нужен kubectl port-forward с локального ПК "
        "(см. раздел 6).",
    )
    add_external_doc_citation(
        doc,
        CARETTA_URL,
        "Caretta's helm chart ships an instance of Grafana with a predefined dashboard using data "
        "published by Caretta.",
        "Helm chart Caretta поставляет экземпляр Grafana с преднастроенным дашбордом на основе "
        "данных, которые публикует Caretta.",
    )

    # 5
    g11.add_heading2(doc, "5. Схема использования имеющихся компонентов", "sec_cu_5")
    add_figure(
        doc,
        IMAGE,
        "Рисунок 1. Radar (Resources) vs Caretta + Grafana (карта трафика)",
    )
    g11.add_normal(doc, "Пояснение к рисунку 1:", consolas=True)
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "Слева — Radar на локальном ПК: режим Topology → Resources показывает Deployment, Service, Pod. "
        "Режим Traffic для вашей установки не активируется.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "В центре — namespace dev: агенты Caretta собирают сетевые связи; Victoria Metrics хранит метрики; "
        "Grafana читает метрики и рисует Node Graph.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "Справа — браузер на localhost:3000 после port-forward к svc/caretta-grafana. "
        "Это основной способ «воспользоваться Caretta» с имеющимися компонентами.",
    )

    # 6
    g11.add_heading2(doc, "6. Пошаговая инструкция: Caretta через Grafana", "sec_cu_6")
    g11.add_normal(doc, "Шаг 1. Терминал Git Bash — port-forward (оставить открытым):")
    g11.add_platform_block(
        doc,
        "Локальный ПК — Windows (Git Bash), терминал 1",
        "export KUBECONFIG=/c/Users/sky/.kube/timeweb-greeting.yaml\n"
        "kubectl port-forward -n dev svc/caretta-grafana 3000:80",
    )
    g11.add_normal(doc, "Успех: строки Forwarding from 127.0.0.1:3000 -> 80. Терминал не закрывать.")
    g11.add_empty(doc)
    g11.add_normal(doc, "Шаг 2. Терминал Git Bash — пароль admin Grafana:")
    g11.add_platform_block(
        doc,
        "Локальный ПК — Windows (Git Bash), терминал 2",
        "export KUBECONFIG=/c/Users/sky/.kube/timeweb-greeting.yaml\n"
        "kubectl get secret caretta-grafana -n dev -o jsonpath='{.data.admin-password}' | base64 -d\n"
        'echo ""',
    )
    g11.add_empty(doc)
    g11.add_normal(doc, "Шаг 3. Браузер на локальном ПК:")
    g11.add_normal(doc, "• URL: http://localhost:3000", consolas=True)
    g11.add_normal(doc, "• Логин: admin", consolas=True)
    g11.add_normal(doc, "• Пароль: значение из шага 2", consolas=True)
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "Шаг 4. В Grafana откройте дашборд Caretta (часто задан как home dashboard в chart). "
        "Панель Node Graph показывает связи client → server по метрике caretta_links_observed.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "Шаг 5. Если карта пустая — сгенерируйте HTTP-трафик к greeting-service (раздел 7), "
        "подождите 15–30 секунд (pollInterval chart) и обновите дашборд.",
    )

    # 7
    g11.add_heading2(doc, "7. Генерация трафика для карты", "sec_cu_7")
    g11.add_normal(doc, "Caretta показывает только реальные соединения. Примеры запросов:")
    g11.add_platform_block(
        doc,
        "Локальный ПК — Windows (Git Bash)",
        "# Через Ingress (подставьте INGRESS_IP при необходимости):\n"
        "curl -H \"Host: greeting-dev.cloud-terra.online\" http://194.87.187.231/api/greeting\n\n"
        "# Или port-forward к приложению:\n"
        "kubectl port-forward -n dev svc/greeting-service 8080:80\n"
        "# В другом терминале:\n"
        "curl -s http://localhost:8080/api/greeting",
    )
    g11.add_normal(
        doc,
        "После нескольких curl на дашборде Caretta должны появиться рёбра между сервисами "
        "(Ingress, greeting-service, DNS, PostgreSQL и т.д.).",
    )

    # 8
    g11.add_heading2(doc, "8. Что использовать Radar", "sec_cu_8")
    g11.add_normal(doc, "Radar по-прежнему полезен для:")
    g11.add_normal(doc, "• Topology → Resources — иерархия объектов;", consolas=True)
    g11.add_normal(doc, "• логи, describe, exec, port-forward к pod;", consolas=True)
    g11.add_normal(doc, "• Helm releases, события кластера.", consolas=True)
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "Radar Traffic (live-потоки на графе) для chart groundcover/caretta сейчас не используйте — "
        "ожидайте operator+CRD или другой источник (Cilium/Hubble, Istio).",
    )

    # 9
    g11.add_heading2(doc, "9. Проверка состояния Caretta", "sec_cu_9")
    g11.add_platform_block(
        doc,
        "Локальный ПК — Windows (Git Bash)",
        "export KUBECONFIG=/c/Users/sky/.kube/timeweb-greeting.yaml\n"
        "helm list -n dev | grep caretta\n"
        "kubectl get pods -n dev | grep caretta\n"
        "kubectl logs -n dev -l app.kubernetes.io/name=caretta --tail=5",
    )
    g11.add_normal(doc, "Признаки нормы:")
    g11.add_normal(doc, "• BPF objects loaded, Kprobe attached — в логах агента;", consolas=True)
    g11.add_normal(doc, "• 2 pod DaemonSet caretta (по числу worker);", consolas=True)
    g11.add_normal(doc, "• caretta-grafana и caretta-vm — Running.", consolas=True)

    # 10
    g11.add_heading2(doc, "10. Итог", "sec_cu_10")
    g11.add_normal(
        doc,
        "1. Caretta установлен и работает (DaemonSet + Victoria Metrics + Grafana в namespace dev).",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "2. Radar Traffic не подключается к этому chart — Radar ждёт Caretta operator и CRD, которых нет.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "3. Grafana не отсутствует — она внутри release caretta, доступ через "
        "kubectl port-forward svc/caretta-grafana 3000:80 и браузер http://localhost:3000.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "4. Имеющиеся компоненты: Radar — для ресурсов и операций; Caretta + Grafana — для карты "
        "сетевого трафика; greeting-service — источник HTTP-нагрузки для наполнения карты.",
    )

    doc.save(OUTPUT)
    print(f"Written: {OUTPUT}")


if __name__ == "__main__":
    main()
