# -*- coding: utf-8 -*-
"""Generate docs/Razdel-15a-gitlab-cicd.docx — Section 15a per Part 2 original structure."""
from __future__ import annotations

import importlib.util
import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches

ROOT = Path(r"D:/Project_infra/greeting-service-infra")
TEMPLATE_CANDIDATES = [
    ROOT / "docs/Razdel-14-helm.docx",
    ROOT / "docs/Razdel-13-kubernetes-manual.docx",
    ROOT / "docs/Traffic-Visibility-Caretta.docx",
    ROOT / "docs/Razdel-flyway-migrations.docx",
]
OUTPUT = ROOT / "docs/Razdel-15a-gitlab-cicd.docx"
BUILD = ROOT / "docs/_build_Razdel-15a-gitlab-cicd.docx"
IMAGE = ROOT / "docs/images/gitlab/gitlab-ci-pipeline.png"
DIAGRAM_SCRIPT = ROOT / "docs/gitlab_ci_diagram_render.py"
CI_FILE = "ci/.gitlab-ci.yml"

G11 = Path(r"C:/Users/sky/AppData/Local/Temp/gen_razdel11a_docx.py")
spec = importlib.util.spec_from_file_location("g11", G11)
g11 = importlib.util.module_from_spec(spec)
assert spec.loader
spec.loader.exec_module(g11)

GITLAB_PIPELINES_URL = "https://docs.gitlab.com/ci/pipelines/"
GITLAB_VARIABLES_URL = "https://docs.gitlab.com/ee/ci/variables/"
GITLAB_CONFIG_PATH_URL = "https://docs.gitlab.com/ee/ci/pipelines/settings.html#specify-a-custom-cicd-configuration-file"

DEVTOOLS_IP = "72.56.249.137"
GITLAB_URL = f"http://{DEVTOOLS_IP}"
PROJECT_PATH = "greeting-group/greeting-service"
HOST_DEV = "greeting-dev.cloud-terra.online"
KUBECONFIG_GB = "export KUBECONFIG=/c/Users/sky/.kube/timeweb-greeting.yaml"
REPO_CD = "cd '/d/Project_infra/greeting-service-infra'"

TOC = [
    ("15а.1. Цель", "sec15a_goal"),
    ("15а.2. Что делается на локальном ПК", "sec15a_local"),
    ("15а.3. Что делается на сервере", "sec15a_server"),
    ("15а.4. Как проверить результат", "sec15a_verify"),
    ("15а.5. Типичные ошибки", "sec15a_errors"),
]


def add_external_doc_citation(doc: Document, url: str, quote_en: str, quote_ru: str) -> None:
    g11.add_empty(doc)
    g11.add_normal(doc, f"Источник: {url}")
    g11.add_empty(doc)
    g11.add_normal(doc, "Цитата:")
    g11.add_platform_block(doc, "Оригинал (English)", quote_en)
    g11.add_empty(doc)
    g11.add_normal(doc, "Перевод:")
    g11.add_platform_block(doc, "Перевод на русский", quote_ru)


def add_toc(doc: Document) -> None:
    g11.add_heading2(doc, "Оглавление", "sec15a_toc")
    for title, anchor in TOC:
        g11.add_hyperlink_paragraph(doc, title, anchor, indent=False)
    g11.add_empty(doc)


def add_gitbash(doc: Document, code: str, *, with_kube: bool = False) -> None:
    prefix = f"export DEVTOOLS_IP={DEVTOOLS_IP}\n{REPO_CD}\n"
    if with_kube:
        prefix += f"{KUBECONFIG_GB}\n"
    g11.add_platform_block(doc, "Локальный ПК — Windows (Git Bash), корень репозитория", prefix + code)


def add_figure(doc: Document, path: Path, caption: str) -> None:
    g11.add_empty(doc)
    g11.add_normal(doc, caption, consolas=True)
    g11.add_empty(doc)
    doc.add_picture(str(path), width=Inches(6.5))
    g11.add_empty(doc)


def ensure_diagram() -> None:
    if not IMAGE.exists():
        subprocess.run([sys.executable, str(DIAGRAM_SCRIPT)], check=True)


def resolve_template() -> Path:
    for path in TEMPLATE_CANDIDATES:
        if path.exists():
            return path
    raise FileNotFoundError("No Word template found in docs/")


def main() -> None:
    ensure_diagram()
    g11._bookmark_id = 0
    shutil.copy2(resolve_template(), BUILD)
    doc = Document(BUILD)
    g11.clear_document_body(doc)

    g11.add_normal(doc, "Раздел 15а. GitLab CI/CD: настройка и запуск pipeline", consolas=True)
    g11.add_empty(doc)
    g11.add_normal(doc, "Версия: 3.0 | 2026-06 | Проект: greeting-service-infra")
    g11.add_empty(doc)
    add_toc(doc)

    # ── 15а.1 Цель ──────────────────────────────────────────────────────────
    g11.add_heading2(doc, "15а.1. Цель", "sec15a_goal")
    g11.add_normal(
        doc,
        "1. Автоматизировать сборку, тестирование, Docker build, Docker push и деплой в Kubernetes "
        "при каждом git push в нужную ветку репозитория GitLab.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "2. В этом разделе рассматривается настройка проекта, CI/CD variables и первого запуска pipeline. "
        "Установка и регистрация GitLab Runner как отдельного компонента рассматривается в Разделе 11a. "
        "Установка GitLab CE — в Разделе 10a.",
    )
    g11.add_empty(doc)
    g11.add_normal(doc, "3. Документация GitLab по pipelines:")
    g11.add_normal(doc, GITLAB_PIPELINES_URL)
    add_external_doc_citation(
        doc,
        GITLAB_PIPELINES_URL,
        "A pipeline is a top-level component for continuous integration, delivery, and deployment.",
        "Pipeline — верхнеуровневый компонент непрерывной интеграции, доставки и развёртывания.",
    )
    g11.add_empty(doc)
    g11.add_normal(doc, "Предварительные условия (другие разделы):")
    g11.add_normal(doc, "• Раздел 10a — GitLab CE, группа greeting-group, проект greeting-service;", consolas=True)
    g11.add_normal(doc, "• Раздел 11a — GitLab Runner (shell executor, тег self-hosted);", consolas=True)
    g11.add_normal(doc, "• Раздел 12 — Docker Registry (docker/docker), Secret в Kubernetes;", consolas=True)
    g11.add_normal(doc, "• Раздел 13–14 — кластер Kubernetes и Helm chart.", consolas=True)

    # ── 15а.2 Локальный ПК ───────────────────────────────────────────────────
    g11.add_heading2(doc, "15а.2. Что делается на локальном ПК", "sec15a_local")

    g11.add_normal(
        doc,
        "1. С локального ПК разработчика через браузер открывается веб-интерфейс GitLab "
        "и создаётся проект с репозиторием для приложения.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        f"2. Открыть адрес {GITLAB_URL}/, убедиться что группа greeting-group и проект "
        f"greeting-service существуют (созданы в Разделе 10a). "
        f"Project → Code → Clone → скопировать Clone URL (HTTP).",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        f"3. После этого на локальном ПК инициализируется Git, добавляется remote GitLab "
        f"и выполняется первый push вместе с конфигурацией {CI_FILE}.",
    )
    g11.add_normal(
        doc,
        "Выполняется на локальном ПК разработчика: macOS (Terminal), Ubuntu (Terminal), "
        "Windows (Git Bash). Назначение: инициализировать локальный Git-репозиторий, "
        "добавить remote GitLab и отправить первый коммит с конфигурацией pipeline.",
    )
    add_gitbash(
        doc,
        "git remote -v\n"
        f"git remote add gitlab {GITLAB_URL}/{PROJECT_PATH}.git 2>/dev/null || true\n"
        "git push -u gitlab master   # или develop — ваша основная ветка",
    )
    g11.add_normal(
        doc,
        "При HTTP push GitLab запросит Username (root) и Password — Personal Access Token "
        "со scope write_repository (см. Раздел 11a.3.5). "
        "GitLab хранит весь монорепозиторий: app/, infra/, ci/, scripts/, docs/.",
    )
    g11.add_empty(doc)

    g11.add_normal(
        doc,
        "4. Далее в интерфейсе GitLab настраиваются CI/CD variables, которые будут использоваться "
        "для работы с Docker Registry, Kubernetes и Helm.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "5. В разделе Settings → CI/CD → Variables добавляются переменные REGISTRY_HOST, "
        "REGISTRY_USER, REGISTRY_PASSWORD, IMAGE_NAME, KUBE_CONFIG_BASE64 и HELM_RELEASE_NAME. "
        "Секретные значения должны быть защищены и скрыты (Masked).",
    )
    g11.add_platform_block(
        doc,
        "Project → Settings → CI/CD → Variables",
        f"REGISTRY_HOST         {DEVTOOLS_IP}:5000\n"
        "REGISTRY_USER         docker\n"
        "REGISTRY_PASSWORD     docker          # Masked\n"
        "IMAGE_NAME            greeting-service\n"
        "HELM_RELEASE_NAME     greeting-service\n"
        "KUBE_CONFIG_BASE64    <base64 kubeconfig>  # Masked\n"
        "REACTIVE_DEMO_IMAGE_NAME      reactive-demo      # опционально\n"
        "REACTIVE_DEMO_HELM_RELEASE    reactive-demo      # опционально",
        yaml_block=True,
    )
    g11.add_normal(
        doc,
        "REGISTRY_USER/PASSWORD — docker/docker из scripts/setup-registry.sh "
        "(не registryuser из устаревших черновиков).",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        f"5.1. Указать путь к файлу конфигурации pipeline (обязательно для монорепозитория):",
    )
    g11.add_normal(doc, "Project → Settings → General → CI/CD → CI/CD configuration file:")
    g11.add_normal(doc, f"   {CI_FILE}", consolas=True)
    add_external_doc_citation(
        doc,
        GITLAB_CONFIG_PATH_URL,
        "You can specify a custom CI/CD configuration file for your project.",
        "Для проекта можно указать пользовательский файл конфигурации CI/CD.",
    )

    g11.add_normal(
        doc,
        "Выполняется на локальном ПК разработчика: macOS (Terminal), Ubuntu (Terminal), "
        "Windows (Git Bash). Назначение: получить kubeconfig в формате base64 "
        "для загрузки в GitLab CI/CD variables.",
    )
    g11.add_normal(doc, "Git Bash (Windows):")
    add_gitbash(doc, "base64 -w 0 /c/Users/sky/.kube/timeweb-greeting.yaml")
    g11.add_normal(doc, "PowerShell (Windows):")
    g11.add_platform_block(
        doc,
        "Локальный ПК — Windows (PowerShell)",
        '[Convert]::ToBase64String([System.IO.File]::ReadAllBytes("$env:USERPROFILE\\.kube\\timeweb-greeting.yaml"))',
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "6. Полученное значение необходимо вставить в переменную KUBE_CONFIG_BASE64 "
        "в настройках проекта GitLab.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "7. После настройки переменных выполняется тестовый запуск pipeline через feature-ветку, "
        "затем полный сценарий запускается через develop или main.",
    )
    g11.add_normal(
        doc,
        "Выполняется на локальном ПК разработчика: macOS (Terminal), Ubuntu (Terminal), "
        "Windows (Git Bash). Назначение: создать feature-ветку и выполнить тестовый push "
        "для запуска сборки и тестов в GitLab CI.",
    )
    add_gitbash(
        doc,
        "git checkout -b feature/test-pipeline\n"
        "# внести небольшое изменение, commit\n"
        "git add .\n"
        "git commit -m \"test: trigger GitLab CI build-and-test\"\n"
        "git push gitlab feature/test-pipeline",
    )
    g11.add_normal(
        doc,
        "Ожидаемый job для feature/*: только build-and-test (без docker и deploy). "
        f"Проверка: {GITLAB_URL}/{PROJECT_PATH}/-/pipelines",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "8. После проверки стартового pipeline изменения переносятся в целевую ветку, "
        "чтобы запустить публикацию образа и деплой в dev-окружение.",
    )
    g11.add_normal(
        doc,
        "Выполняется на локальном ПК разработчика: macOS (Terminal), Ubuntu (Terminal), "
        "Windows (Git Bash). Назначение: выполнить merge в целевую ветку и запустить "
        "полный GitLab pipeline для dev-окружения.",
    )
    add_gitbash(
        doc,
        "git checkout develop\n"
        "git merge feature/test-pipeline\n"
        "git push gitlab develop",
    )
    g11.add_normal(
        doc,
        "Ожидаемая цепочка job для develop: build-and-test → build-and-push-docker → deploy-dev.",
    )
    g11.add_empty(doc)
    g11.add_normal(doc, "Схема pipeline (справочно):")
    add_figure(doc, IMAGE, "Рисунок 1. GitLab CI/CD pipeline greeting-service-infra")
    g11.add_normal(
        doc,
        f"Файл {CI_FILE}: stages build → docker → deploy; Runner shell на devtools (Раздел 11a); "
        "IMAGE_TAG = CI_COMMIT_SHORT_SHA.",
    )

    # ── 15а.3 Сервер ──────────────────────────────────────────────────────────
    g11.add_heading2(doc, "15а.3. Что делается на сервере", "sec15a_server")
    g11.add_normal(
        doc,
        "1. В этом разделе прямые команды на devtools-сервере обычно не выполняются, "
        "потому что запуск GitLab CI/CD инициируется из локального Git-клиента и веб-интерфейса GitLab.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        f"2. На стороне облака GitLab хранит проект и конфигурацию {CI_FILE}, "
        "а выполнение pipeline обеспечивается GitLab Runner (Раздел 11a, shell executor, тег self-hosted). "
        "Job выполняются на devtools от пользователя gitlab-runner.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "3. Docker Registry на devtools-сервере принимает собранные образы, "
        "а Kubernetes-кластер принимает изменения через kubectl и helm внутри job GitLab CI.",
    )
    g11.add_empty(doc)
    g11.add_normal(doc, "Что Runner выполняет на devtools (без ручного SSH):")
    g11.add_platform_block(
        doc,
        "Job на devtools (shell executor)",
        "cd app && ./gradlew clean test bootJar\n"
        f"docker build -t {DEVTOOLS_IP}:5000/greeting-service:$CI_COMMIT_SHORT_SHA app/\n"
        f"docker push {DEVTOOLS_IP}:5000/greeting-service:$CI_COMMIT_SHORT_SHA\n"
        "helm upgrade --install greeting-service infra/helm/greeting-service ...\n"
        "kubectl rollout status deployment/greeting-service -n dev",
        yaml_block=True,
    )

    # ── 15а.4 Проверка ────────────────────────────────────────────────────────
    g11.add_heading2(doc, "15а.4. Как проверить результат", "sec15a_verify")
    g11.add_normal(
        doc,
        "1. После запуска pipeline необходимо проверить статус job в GitLab UI "
        "и убедиться, что в Kubernetes обновилось приложение.",
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "2. В интерфейсе GitLab в разделе Build → Pipelines или Build → Jobs "
        "все job нужного запуска должны завершиться успешно.",
    )
    g11.add_normal(
        doc,
        "Выполняется на локальном ПК разработчика: macOS (Terminal), Ubuntu (Terminal), "
        "Windows (Git Bash). Назначение: проверить состояние pod после деплоя и убедиться, "
        "что в dev развернулась новая версия приложения.",
    )
    add_gitbash(
        doc,
        f'curl -u docker:docker "http://{DEVTOOLS_IP}:5000/v2/greeting-service/tags/list"\n'
        "kubectl get pods,deployment,ingress -n dev\n"
        "helm list -n dev\n"
        f'curl -s "http://{HOST_DEV}/api/greeting"',
        with_kube=True,
    )
    g11.add_empty(doc)
    g11.add_normal(
        doc,
        "3. После успешного pipeline pod должны быть обновлены, "
        "а новая версия образа должна использоваться в namespace dev.",
    )
    g11.add_normal(
        doc,
        "Успех: тег образа в registry совпадает с CI_COMMIT_SHORT_SHA коммита; "
        "curl /api/greeting возвращает JSON с полем message.",
    )

    # ── 15а.5 Ошибки ──────────────────────────────────────────────────────────
    g11.add_heading2(doc, "15а.5. Типичные ошибки", "sec15a_errors")

    g11.add_normal(doc, "Ошибка: Pipeline не стартует или зависает в состоянии ожидания исполнителя.", consolas=True)
    g11.add_normal(
        doc,
        "Причина: Для проекта недоступна корректная среда исполнения GitLab CI "
        "либо конфигурация job не соответствует доступной среде запуска.",
    )
    g11.add_normal(
        doc,
        "Исправление: Проверить Runner в GitLab UI (Project → CI/CD → Runners — Online). "
        "На devtools: sudo gitlab-runner status. См. Раздел 11a.",
    )
    g11.add_empty(doc)

    g11.add_normal(
        doc,
        "Ошибка: В job появляется сообщение docker: command not found "
        "или permission denied while trying to connect to the Docker daemon.",
        consolas=True,
    )
    g11.add_normal(
        doc,
        "Причина: Используется окружение без Docker либо среда выполнения job "
        "не имеет прав на работу с Docker.",
    )
    g11.add_normal(
        doc,
        "Исправление: Для shell executor — sudo usermod -aG docker gitlab-runner; "
        "sudo gitlab-runner restart. Проверка: sudo -u gitlab-runner docker ps.",
    )
    g11.add_empty(doc)

    g11.add_normal(doc, "Ошибка: Unable to create pipeline — CI configuration not found.", consolas=True)
    g11.add_normal(doc, f"Причина: не указан путь {CI_FILE} в Settings → CI/CD configuration file.", consolas=True)
    g11.add_empty(doc)

    g11.add_normal(doc, "Ошибка: permission denied при ./gradlew.", consolas=True)
    g11.add_normal(doc, "Исправление: chmod +x app/gradlew в script job (уже есть в ci/.gitlab-ci.yml).", consolas=True)
    g11.add_empty(doc)

    g11.add_normal(doc, "Ошибка: После деплоя pod завершается с ошибкой exec format error.", consolas=True)
    g11.add_normal(
        doc,
        "Причина: Docker-образ был собран для другой архитектуры, например arm64, "
        "тогда как узлы Kubernetes используют amd64.",
    )
    g11.add_normal(
        doc,
        "Исправление: Явно указать целевую платформу в Dockerfile и повторно выполнить сборку через GitLab pipeline.",
    )
    g11.add_normal(
        doc,
        "Выполняется на локальном ПК разработчика, в Dockerfile приложения. "
        "Назначение: зафиксировать целевую платформу образа для совместимости с узлами Kubernetes.",
    )
    g11.add_platform_block(
        doc,
        "app/Dockerfile (фрагмент)",
        "# Добавить при сборке на arm64-хосте:\n"
        "# docker build --platform linux/amd64 ...\n"
        "# или в Dockerfile: FROM --platform=linux/amd64 eclipse-temurin:21-jre",
        yaml_block=True,
    )
    g11.add_empty(doc)

    g11.add_normal(doc, "Ошибка: ImagePullBackOff после deploy-dev.", consolas=True)
    g11.add_normal(doc, "Причина: образ с тегом SHA не попал в registry (job docker failed).", consolas=True)
    g11.add_normal(
        doc,
        f'Исправление: логи build-and-push-docker; '
        f'curl -u docker:docker "http://{DEVTOOLS_IP}:5000/v2/greeting-service/tags/list"',
    )

    doc.save(BUILD)
    try:
        if OUTPUT.exists():
            OUTPUT.unlink()
        BUILD.replace(OUTPUT)
        saved = OUTPUT
        if BUILD.exists():
            BUILD.unlink()
    except PermissionError:
        saved = BUILD
    print(f"Written: {saved}")


if __name__ == "__main__":
    main()
