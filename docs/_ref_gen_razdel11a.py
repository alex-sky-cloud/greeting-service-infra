# -*- coding: utf-8 -*-
"""Generate docs/Razdel-11a-gitlab-runner.docx from original Section 11a."""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(r"D:/!_Проекты инфраструктуры/greeting-service-infra")
TEMPLATE = ROOT / "docs/Razdel-9 (2).docx"
OUTPUT = ROOT / "docs/Razdel-11a-gitlab-runner.docx"

SHADE_FILL = "E7E6E6"
DASH = "-----------------------------------------------------------------------"
CAPTION_BLUE = RGBColor(0x00, 0xB0, 0xF0)

BASH_KW = {
    "if", "then", "else", "fi", "for", "do", "done", "in", "return", "set",
    "export", "source", "echo", "cd", "chmod", "mkdir", "sudo", "curl", "ssh",
    "bash", "apt-get", "wget", "tar", "tee", "ln", "rm", "docker", "docker-compose",
    "gitlab-runner", "gitlab-ctl", "journalctl", "usermod", "systemctl",
}
COL = {
    "text": RGBColor(0x00, 0x00, 0x00),
    "comment": RGBColor(0xA0, 0xA1, 0xA7),
    "keyword": RGBColor(0x4C, 0x62, 0xAF),
    "string": RGBColor(0x66, 0x0E, 0x7A),
    "variable": RGBColor(0xAF, 0x27, 0xAD),
    "command": RGBColor(0x19, 0x95, 0xA0),
}

TOC = [
    ("11a.1. Цель", "sec11a_goal"),
    ("11a.2. Предварительные условия", "sec11a_prereq"),
    ("11a.3. Что делается на локальном ПК", "sec11a_local"),
    ("11a.3.1. IP devtools и проверка SSH", "sec11a_local_ssh"),
    ("11a.3.2. Создание Runner в GitLab", "sec11a_local_ui"),
    ("11a.4. Что делается на devtools-сервере", "sec11a_server"),
    ("11a.4.1. Установка пакета GitLab Runner", "sec11a_server_install"),
    ("11a.4.2. Регистрация Runner", "sec11a_server_register"),
    ("11a.5. Пример тега в .gitlab-ci.yml", "sec11a_gitlab_ci"),
    ("11a.6. Как проверить результат", "sec11a_verify"),
    ("11a.7. Типичные ошибки", "sec11a_errors"),
]

_bookmark_id = 0


def next_bookmark_id() -> int:
    global _bookmark_id
    _bookmark_id += 1
    return _bookmark_id


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def is_toc_subsection(title: str) -> bool:
    return bool(re.match(r"11a\.\d+\.\d+", title))


def add_hyperlink_paragraph(doc: Document, text: str, anchor: str, *, indent: bool = False) -> None:
    p = doc.add_paragraph()
    p_pr = p._element.get_or_add_pPr()
    if indent:
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "397")
        p_pr.append(ind)
    para_r_pr = OxmlElement("w:rPr")
    para_fonts = OxmlElement("w:rFonts")
    para_fonts.set(qn("w:ascii"), "Consolas")
    para_fonts.set(qn("w:hAnsi"), "Consolas")
    para_r_pr.append(para_fonts)
    p_pr.append(para_r_pr)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Consolas")
    fonts.set(qn("w:hAnsi"), "Consolas")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_pr.append(fonts)
    run_pr.append(color)
    run_pr.append(underline)
    t = OxmlElement("w:t")
    t.text = text
    run.append(run_pr)
    run.append(t)
    hyperlink.append(run)
    p._p.append(hyperlink)


def clear_document_body(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_shading(paragraph, fill: str = SHADE_FILL) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_run(paragraph, text: str, *, font="Consolas", size=10, color: RGBColor | None = None, bold=False):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), font)
    r_fonts.set(qn("w:hAnsi"), font)
    r_pr.append(r_fonts)
    return run


def add_empty(doc: Document):
    doc.add_paragraph("")


def add_normal(doc: Document, text: str, *, consolas=False):
    p = doc.add_paragraph()
    if consolas:
        add_run(p, text, color=CAPTION_BLUE)
    else:
        p.add_run(text)
    return p


def add_heading2(doc: Document, text: str, bookmark: str):
    p = doc.add_paragraph(text, style="Heading 2")
    for r in p.runs:
        r.font.name = "Consolas"
    add_bookmark(p, bookmark, next_bookmark_id())
    return p


def add_heading3(doc: Document, text: str, bookmark: str):
    p = doc.add_paragraph(text, style="Heading 3")
    for r in p.runs:
        r.font.name = "Consolas"
    add_bookmark(p, bookmark, next_bookmark_id())
    return p


def add_toc(doc: Document) -> None:
    add_heading2(doc, "Оглавление", "sec11a_toc")
    for title, anchor in TOC:
        add_hyperlink_paragraph(doc, title, anchor, indent=is_toc_subsection(title))
    add_empty(doc)


def tokenize_bash_line(line: str) -> list[tuple[str, str]]:
    tokens: list[tuple[str, str]] = []
    i = 0
    first = True
    while i < len(line):
        if line[i].isspace():
            j = i
            while j < len(line) and line[j].isspace():
                j += 1
            tokens.append(("text", line[i:j]))
            i = j
            continue
        if line.startswith("#", i):
            tokens.append(("comment", line[i:]))
            break
        if line[i] in "\"'":
            q = line[i]
            j = i + 1
            while j < len(line):
                if line[j] == q and line[j - 1] != "\\":
                    j += 1
                    break
                j += 1
            tokens.append(("string", line[i:j]))
            i = j
            continue
        if line[i] == "$" and i + 1 < len(line) and line[i + 1] == "{":
            j = i + 2
            depth = 1
            while j < len(line) and depth:
                if line[j] == "{":
                    depth += 1
                elif line[j] == "}":
                    depth -= 1
                j += 1
            tokens.append(("variable", line[i:j]))
            i = j
            continue
        j = i
        while j < len(line) and not line[j].isspace() and line[j] not in "\"'$":
            j += 1
        if j == i:
            tokens.append(("text", line[i]))
            i += 1
            first = False
            continue
        word = line[i:j]
        if first and word in BASH_KW:
            kind = "keyword"
        elif first:
            kind = "command"
        else:
            kind = "text"
        tokens.append((kind, word))
        first = False
        i = j
    return tokens


def append_bash_line(paragraph, line: str, *, first_line: bool):
    if not first_line:
        paragraph.add_run().add_break()
    if not line:
        return
    for kind, token in tokenize_bash_line(line):
        add_run(paragraph, token, color=COL.get(kind, COL["text"]))


def add_code_block(doc: Document, code: str, *, bash=True):
    add_empty(doc)
    p = doc.add_paragraph()
    set_shading(p)
    body = code.strip("\n")
    full = f"{DASH}\n{body}\n{DASH}"
    if bash:
        for idx, line in enumerate(full.splitlines()):
            append_bash_line(p, line, first_line=(idx == 0))
    else:
        for idx, line in enumerate(full.splitlines()):
            if idx:
                p.add_run().add_break()
            add_run(p, line)
    add_empty(doc)


def add_platform_block(doc: Document, platform: str, code: str, *, yaml_block=False):
    add_normal(doc, platform, consolas=True)
    add_code_block(doc, code, bash=not yaml_block)


def add_command_block(doc: Document, code: str, explanation: str | None = None):
    add_normal(doc, "Команда:", consolas=True)
    add_code_block(doc, code)
    if explanation:
        add_normal(doc, explanation)


def main() -> None:
    global _bookmark_id
    _bookmark_id = 0

    shutil.copy2(TEMPLATE, OUTPUT)
    doc = Document(OUTPUT)
    clear_document_body(doc)

    add_normal(doc, "Раздел 11a. Инструкция по GitLab Runner (self-hosted)", consolas=True)
    add_empty(doc)
    add_toc(doc)

    add_heading2(doc, "11a.1. Цель", "sec11a_goal")
    add_normal(
        doc,
        "1. Настроить self-hosted GitLab Runner для выполнения pipeline-задач проекта на devtools-сервере "
        "в облаке. Такой Runner позволяет запускать сборку, тесты и публикацию Docker-образов внутри "
        "собственной инфраструктуры.",
    )
    add_empty(doc)
    add_normal(
        doc,
        "2. В разделе рассматривается только установка и регистрация GitLab Runner. Настройка GitLab, "
        "Docker Registry и схема деплоя рассматривается в других разделах.",
    )

    add_heading2(doc, "11a.2. Предварительные условия", "sec11a_prereq")
    add_normal(doc, "1. Выполнен Раздел 10a: GitLab CE установлен, группа greeting-group и проект greeting-service созданы.")
    add_empty(doc)
    add_normal(doc, "2. Docker Registry развёрнут (порт 5000, учётные данные docker / docker).")
    add_empty(doc)
    add_normal(doc, "3. Роли терминалов:")
    add_normal(doc, "- terraform output — WSL Ubuntu;", consolas=True)
    add_normal(doc, "- ssh, curl — Git Bash (Windows);", consolas=True)
    add_empty(doc)
    add_normal(doc, "4. SSH-ключ: C:\\Users\\sky\\.ssh\\id_ed25519; пользователь на devtools — root (не ubuntu).")

    add_heading2(doc, "11a.3. Что делается на локальном ПК", "sec11a_local")
    add_normal(
        doc,
        "1. На локальном ПК определяется публичный IP devtools-сервера и проверяется SSH-доступ к нему.",
    )
    add_empty(doc)
    add_normal(doc, "2. В браузере в GitLab создаётся новый Runner для проекта или группы.")
    add_empty(doc)
    add_normal(doc, "3. Сохраняется authentication token — он потребуется на devtools-сервере при регистрации Runner.")

    add_heading3(doc, "11a.3.1. IP devtools и проверка SSH", "sec11a_local_ssh")
    add_normal(doc, "Получить IP (WSL Ubuntu):")
    add_platform_block(
        doc,
        "Локальный ПК — WSL Ubuntu",
        "cd '/mnt/d/!_Проекты инфраструктуры/greeting-service-infra/infra/terraform'\n"
        "terraform output -raw devtools_public_ip",
    )
    add_normal(doc, "Сохранить IP и проверить SSH (Git Bash):")
    add_platform_block(
        doc,
        "Локальный ПК — Windows (Git Bash), корень репозитория",
        "cd '/d/!_Проекты инфраструктуры/greeting-service-infra'\n"
        "DEVTOOLS_IP=72.56.249.137   # подставьте terraform output\n"
        "echo \"Devtools IP: ${DEVTOOLS_IP}\"\n"
        "ssh -i /c/Users/sky/.ssh/id_ed25519 root@${DEVTOOLS_IP} \"echo connected\"",
    )
    add_normal(doc, "Успех: строка connected без запроса пароля.")

    add_heading3(doc, "11a.3.2. Создание Runner в GitLab", "sec11a_local_ui")
    add_normal(doc, "1. Открыть GitLab: http://72.56.249.137/users/sign_in (подставьте свой DEVTOOLS_IP).")
    add_empty(doc)
    add_normal(
        doc,
        "2. Перейти в Project → Settings → CI/CD → Runners "
        "либо в Group → Build → Runners — в зависимости от выбранной области регистрации.",
    )
    add_empty(doc)
    add_normal(doc, "3. Создать новый Runner (New project runner / Create group runner).")
    add_empty(doc)
    add_normal(doc, "4. При создании Runner сохранить authentication token (glrt-...).")
    add_empty(doc)
    add_normal(doc, "5. Токен понадобится в п. 11a.4.2 на devtools-сервере в параметре --token.")

    add_heading2(doc, "11a.4. Что делается на devtools-сервере", "sec11a_server")
    add_normal(
        doc,
        "1. На devtools-сервере устанавливается пакет GitLab Runner и выполняется регистрация Runner "
        "для подключения к GitLab.",
    )

    add_heading3(doc, "11a.4.1. Установка пакета GitLab Runner", "sec11a_server_install")
    add_normal(doc, "Подключиться к devtools (Git Bash):")
    add_platform_block(
        doc,
        "Локальный ПК — Windows (Git Bash)",
        "ssh -i /c/Users/sky/.ssh/id_ed25519 root@${DEVTOOLS_IP}",
    )
    add_normal(doc, "На devtools-сервере — установить пакет из официального репозитория:")
    add_platform_block(
        doc,
        "На devtools-сервере (после SSH)",
        "curl -fsSL https://packages.gitlab.com/install/repositories/runner/gitlab-runner/script.deb.sh | sudo bash\n"
        "sudo apt-get install -y gitlab-runner",
    )

    add_heading3(doc, "11a.4.2. Регистрация Runner", "sec11a_server_register")
    add_normal(doc, "На devtools-сервере. Сначала задать IP — переменная из Git Bash сюда не переносится:")
    add_platform_block(
        doc,
        "На devtools-сервере (после SSH)",
        "DEVTOOLS_IP=72.56.249.137   # подставьте terraform output\n\n"
        "sudo gitlab-runner register \\\n"
        "  --non-interactive \\\n"
        "  --url \"http://${DEVTOOLS_IP}/\" \\\n"
        "  --token \"<runner-authentication-token>\" \\\n"
        "  --executor \"shell\" \\\n"
        "  --description \"devtools-runner\" \\\n"
        "  --tag-list \"self-hosted,devtools\" \\\n"
        "  --run-untagged=\"false\" \\\n"
        "  --locked=\"false\"",
    )
    add_normal(doc, "В параметре --token указать authentication token, полученный в GitLab (п. 11a.3.2).")
    add_empty(doc)
    add_normal(
        doc,
        "При использовании shell executor команды pipeline выполняются непосредственно на devtools-сервере "
        "от имени пользователя gitlab-runner.",
    )

    add_heading2(doc, "11a.5. Пример тега в .gitlab-ci.yml", "sec11a_gitlab_ci")
    add_normal(
        doc,
        "1. Для запуска pipeline на зарегистрированном Runner необходимо указать соответствующий тег "
        "в файле .gitlab-ci.yml.",
    )
    add_empty(doc)
    add_normal(doc, "2. Пример фрагмента (корень репозитория greeting-service-infra):")
    add_platform_block(
        doc,
        "Локальный ПК — Windows (Git Bash), корень репозитория",
        "build-and-push:\n"
        "  stage: build\n"
        "  tags:\n"
        "    - self-hosted\n"
        "  script:\n"
        "    - docker version\n"
        "    - docker build -t ${DEVTOOLS_IP}:5000/greeting-service:$CI_COMMIT_SHORT_SHA .\n"
        "    - docker push ${DEVTOOLS_IP}:5000/greeting-service:$CI_COMMIT_SHORT_SHA",
        yaml_block=True,
    )

    add_heading2(doc, "11a.6. Как проверить результат", "sec11a_verify")
    add_normal(
        doc,
        "1. Проверка выполняется на devtools-сервере и в интерфейсе GitLab: состояние службы Runner "
        "и успешная регистрация в GitLab.",
    )
    add_platform_block(
        doc,
        "Локальный ПК — Windows (Git Bash)",
        "ssh -i /c/Users/sky/.ssh/id_ed25519 root@${DEVTOOLS_IP} 'sudo gitlab-runner status'\n"
        "ssh -i /c/Users/sky/.ssh/id_ed25519 root@${DEVTOOLS_IP} 'sudo gitlab-runner verify'\n"
        "ssh -i /c/Users/sky/.ssh/id_ed25519 root@${DEVTOOLS_IP} "
        "'sudo journalctl -u gitlab-runner -n 30 --no-pager'",
    )
    add_normal(doc, "2. Команда gitlab-runner verify должна подтвердить доступность зарегистрированного Runner.")
    add_empty(doc)
    add_normal(
        doc,
        "3. В интерфейсе GitLab Runner отображается в статусе Online. "
        "После запуска тестового pipeline задание с тегом self-hosted должно перейти на этот Runner.",
    )

    add_heading2(doc, "11a.7. Типичные ошибки", "sec11a_errors")
    add_normal(doc, "1. Ошибка: Runner не появляется в GitLab или отображается как Offline.")
    add_normal(doc, "2. Причина: использован неверный URL GitLab либо указан неправильный authentication token при регистрации.")
    add_normal(doc, "3. Исправление: повторно проверить адрес GitLab, создать новый token в интерфейсе и заново выполнить регистрацию Runner.")
    add_empty(doc)
    add_normal(doc, "4. Ошибка: в job появляется сообщение docker: command not found или отказ в доступе к Docker.")
    add_normal(doc, "5. Причина: на devtools-сервере отсутствует Docker либо пользователь gitlab-runner не имеет права работать с Docker-сокетом.")
    add_normal(doc, "6. Исправление: установить Docker, добавить пользователя gitlab-runner в группу docker и перезапустить службу GitLab Runner.")
    add_command_block(
        doc,
        "sudo usermod -aG docker gitlab-runner\n"
        "sudo systemctl restart gitlab-runner",
        "Выполняется на devtools-сервере после SSH.",
    )

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
