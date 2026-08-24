# -*- coding: utf-8 -*-
"""Общий стиль Word-документов (как в Части 2 / Раздел 11a)."""
from __future__ import annotations

import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

SHADE_FILL = "E7E6E6"
DASH = "-----------------------------------------------------------------------"
CAPTION_BLUE = RGBColor(0x00, 0xB0, 0xF0)

BASH_KW = {
    "if", "then", "else", "elif", "fi", "for", "do", "done", "in", "return", "set",
    "export", "source", "echo", "cd", "chmod", "mkdir", "sudo", "curl", "ssh",
    "bash", "apt-get", "apt", "wget", "tar", "tee", "ln", "rm", "docker",
    "docker-compose", "gitlab-runner", "gitlab-ctl", "journalctl", "usermod",
    "systemctl", "helm", "kubectl", "scp", "sed", "cat", "htpasswd", "ufw",
    "openssl", "mc", "kubectl", "base64", "tr", "true", "false",
}

COL = {
    "text": RGBColor(0x00, 0x00, 0x00),
    "comment": RGBColor(0xA0, 0xA1, 0xA7),
    "keyword": RGBColor(0x4C, 0x62, 0xAF),
    "string": RGBColor(0x66, 0x0E, 0x7A),
    "variable": RGBColor(0xAF, 0x27, 0xAD),
    "command": RGBColor(0x19, 0x95, 0xA0),
}

_bookmark_id = 0


def reset_bookmarks() -> None:
    global _bookmark_id
    _bookmark_id = 0


def next_bookmark_id() -> int:
    global _bookmark_id
    _bookmark_id += 1
    return _bookmark_id


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_hyperlink_paragraph(doc: Document, text: str, anchor: str, *, indent: bool = False) -> None:
    p = doc.add_paragraph()
    p_pr = p._element.get_or_add_pPr()
    if indent:
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "397")
        p_pr.append(ind)
    para_r_pr = OxmlElement("w:rPr")
    para_fonts = OxmlElement("w:rFonts")
    para_fonts.set(qn("w:ascii"), "Consolas")
    para_fonts.set(qn("w:hAnsi"), "Consolas")
    para_r_pr.append(para_fonts)
    p_pr.append(para_r_pr)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Consolas")
    fonts.set(qn("w:hAnsi"), "Consolas")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_pr.append(fonts)
    run_pr.append(color)
    run_pr.append(underline)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(run_pr)
    run.append(t)
    hyperlink.append(run)
    p._p.append(hyperlink)


def clear_document_body(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_shading(paragraph, fill: str = SHADE_FILL) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_run(paragraph, text: str, *, font="Consolas", size=10, color: RGBColor | None = None, bold=False):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), font)
    r_fonts.set(qn("w:hAnsi"), font)
    r_pr.append(r_fonts)
    return run


def add_empty(doc: Document):
    doc.add_paragraph("")


def add_normal(doc: Document, text: str, *, consolas=False):
    p = doc.add_paragraph()
    if consolas:
        add_run(p, text, color=CAPTION_BLUE)
    else:
        p.add_run(text)
    return p


def add_heading2(doc: Document, text: str, bookmark: str):
    p = doc.add_paragraph(text, style="Heading 2")
    for r in p.runs:
        r.font.name = "Consolas"
    add_bookmark(p, bookmark, next_bookmark_id())
    return p


def add_heading3(doc: Document, text: str, bookmark: str):
    p = doc.add_paragraph(text, style="Heading 3")
    for r in p.runs:
        r.font.name = "Consolas"
    add_bookmark(p, bookmark, next_bookmark_id())
    return p


def tokenize_bash_line(line: str) -> list[tuple[str, str]]:
    tokens: list[tuple[str, str]] = []
    i = 0
    first = True
    while i < len(line):
        if line[i].isspace():
            j = i
            while j < len(line) and line[j].isspace():
                j += 1
            tokens.append(("text", line[i:j]))
            i = j
            continue
        if line.startswith("#", i):
            tokens.append(("comment", line[i:]))
            break
        if line[i] in "\"'":
            q = line[i]
            j = i + 1
            while j < len(line):
                if line[j] == q and line[j - 1] != "\\":
                    j += 1
                    break
                j += 1
            tokens.append(("string", line[i:j]))
            i = j
            continue
        if line[i] == "$" and i + 1 < len(line) and line[i + 1] == "{":
            j = i + 2
            depth = 1
            while j < len(line) and depth:
                if line[j] == "{":
                    depth += 1
                elif line[j] == "}":
                    depth -= 1
                j += 1
            tokens.append(("variable", line[i:j]))
            i = j
            continue
        j = i
        while j < len(line) and not line[j].isspace() and line[j] not in "\"'$":
            j += 1
        if j == i:
            tokens.append(("text", line[i]))
            i += 1
            first = False
            continue
        word = line[i:j]
        if first and word in BASH_KW:
            kind = "keyword"
        elif first:
            kind = "command"
        else:
            kind = "text"
        tokens.append((kind, word))
        first = False
        i = j
    return tokens


def append_bash_line(paragraph, line: str, *, first_line: bool):
    if not first_line:
        paragraph.add_run().add_break()
    if not line:
        return
    for kind, token in tokenize_bash_line(line):
        add_run(paragraph, token, color=COL.get(kind, COL["text"]))


def add_code_block(doc: Document, code: str, *, bash=True):
    add_empty(doc)
    p = doc.add_paragraph()
    set_shading(p)
    body = code.strip("\n")
    full = f"{DASH}\n{body}\n{DASH}"
    if bash:
        for idx, line in enumerate(full.splitlines()):
            append_bash_line(p, line, first_line=(idx == 0))
    else:
        for idx, line in enumerate(full.splitlines()):
            if idx:
                p.add_run().add_break()
            add_run(p, line)
    add_empty(doc)


def add_platform_block(doc: Document, platform: str, code: str, *, yaml_block=False):
    add_normal(doc, platform, consolas=True)
    add_code_block(doc, code, bash=not yaml_block)


def add_command_block(doc: Document, code: str, explanation: str | None = None):
    add_normal(doc, "Команда:", consolas=True)
    add_code_block(doc, code)
    if explanation:
        add_normal(doc, explanation)


def add_citation(doc: Document, url: str, quote_en: str, quote_ru: str) -> None:
    add_empty(doc)
    add_normal(doc, f"Источник: {url}")
    add_empty(doc)
    add_normal(doc, "Цитата:")
    add_platform_block(doc, "Оригинал (English)", quote_en)
    add_empty(doc)
    add_normal(doc, "Перевод:")
    add_platform_block(doc, "Перевод на русский", quote_ru)


def is_toc_subsection(title: str) -> bool:
    return bool(re.search(r"\d+\.\d+\.\d+", title))
