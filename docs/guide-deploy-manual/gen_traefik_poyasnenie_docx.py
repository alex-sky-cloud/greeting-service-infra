# -*- coding: utf-8 -*-
"""Word-пояснение: зачем Traefik на отдельном VPS (к схеме 1.1)."""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches

HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(HERE))
import docx_style as sty  # noqa: E402

ROOT = Path(r"D:/Project_infra/greeting-service-infra")
OUTPUT = ROOT / "docs" / "Traefik-zachem-otdelnyi-klaster.docx"
BUILD = HERE / "_build_traefik_poyasnenie.docx"

TEMPLATE_CANDIDATES = [
    ROOT / "docs/guide-deploy-terraform/Razdel-15a-gitlab-cicd.docx",
    ROOT / "docs/guide-deploy-terraform/Razdel-flyway-migrations.docx",
    HERE / "greeting-service-guide - Часть 2M - CI-CD вручную (GitLab k3s Traefik).docx",
]

SCREENSHOT_CANDIDATES = [
    Path(
        r"C:/Users/sky/.cursor/projects/d-Project-infra-greeting-service-infra"
        r"/assets/c__Users_sky_AppData_Roaming_Cursor_User_workspaceStorage"
        r"_371f2630de332b8e39996771b99e1f6e_images_Screenshot_291-8e2e33a4-d3d2-4e50-92a1-51e84f1178f8.png"
    ),
    HERE / "images" / "architecture-traffic.png",
]
IMG_SCHEME = HERE / "images" / "scheme-1-1-chto-stroim.png"

TOC = [
    ("1. Что делает Traefik", "trf_s01_what"),
    ("2. Как это видно на схеме 1.1", "trf_s02_scheme"),
    ("3. Зачем он здесь вообще", "trf_s03_why"),
    ("4. Почему на отдельном VPS", "trf_s04_vps"),
    ("5. Почему именно Traefik", "trf_s05_why_traefik"),
    ("6. Одна фраза вместо запутанной", "trf_s06_one_phrase"),
]


def resolve_template() -> Path:
    for p in TEMPLATE_CANDIDATES:
        if p.exists():
            return p
    raise FileNotFoundError("Не найден Word-шаблон со стилями Heading")


def add_toc(doc: Document) -> None:
    sty.add_heading2(doc, "Оглавление", "trf_toc")
    for title, anchor in TOC:
        sty.add_hyperlink_paragraph(doc, title, anchor, indent=False)
    sty.add_empty(doc)


def add_figure(doc: Document, path: Path, caption: str) -> None:
    sty.add_empty(doc)
    sty.add_normal(doc, caption, consolas=True)
    sty.add_empty(doc)
    doc.add_picture(str(path), width=Inches(6.7))
    sty.add_empty(doc)


def ensure_scheme_image() -> Path:
    IMG_SCHEME.parent.mkdir(parents=True, exist_ok=True)
    if IMG_SCHEME.exists():
        return IMG_SCHEME
    for src in SCREENSHOT_CANDIDATES:
        if src.exists():
            shutil.copy2(src, IMG_SCHEME)
            return IMG_SCHEME
    raise FileNotFoundError(f"Нет схемы 1.1: {IMG_SCHEME}")


def build_document(doc: Document) -> None:
    sty.add_normal(doc, "Пояснение к схеме 1.1", consolas=True)
    sty.add_normal(doc, "Зачем Traefik на отдельном VPS", consolas=True)
    sty.add_empty(doc)
    sty.add_normal(doc, "К разделу «1.1. Что строим» гайда Часть 2M (CI/CD вручную).")
    sty.add_normal(doc, "Версия: 1.0 | 2026-08 | Без Terraform, без managed Ingress облака")
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "Фраза «внешний HTTP(S)-трафик принимает отдельный кластер Traefik на своих VPS» "
        "без роли блока выглядит лишней. Этот документ разбирает схему и отвечает: "
        "что делает Traefik, зачем он, почему отдельно от k3s и GitLab.",
    )
    sty.add_empty(doc)
    add_toc(doc)

    sty.add_heading2(doc, "1. Что делает Traefik", "trf_s01_what")
    sty.add_normal(
        doc,
        "Traefik — входная дверь из интернета. Он принимает HTTPS «с улицы» и отводит "
        "запрос в нужный сервис внутри Kubernetes.",
    )
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "Официально это application proxy (прокси приложений / edge-маршрутизатор): "
        "Traefik принимает запросы от имени вашей системы, понимает, какой компонент "
        "должен их обработать, и направляет туда.",
    )
    sty.add_citation(
        doc,
        "https://doc.traefik.io/traefik/",
        "Here’s how it works—Traefik receives requests on behalf of your system, "
        "identifies which components are responsible for handling them, and routes them securely.",
        "Как это работает: Traefik принимает запросы от имени вашей системы, определяет, "
        "какие компоненты должны их обработать, и безопасно маршрутизирует их.",
    )
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "Пользователь не стучится ни в GitLab, ни напрямую в Pod. Он стучится в Traefik.",
    )

    sty.add_heading2(doc, "2. Как это видно на схеме 1.1", "trf_s02_scheme")
    sty.add_normal(
        doc,
        "На схеме Traefik стоит не внутри k3s и не рядом с GitLab. Он стоит на пути "
        "пользователя. Это другая линия, чем git push.",
    )
    add_figure(
        doc,
        IMG_SCHEME,
        "Рисунок. Схема 1.1 «Что строим»: две линии — CI/CD сверху, HTTPS снизу",
    )
    sty.add_normal(doc, "По схеме путь пользователя такой:")
    sty.add_normal(doc, "1. Пользователь (глобус) шлёт HTTPS.")
    sty.add_normal(doc, "2. Стрелка идёт только в фиолетовый блок Traefik (:80 / :443).")
    sty.add_normal(doc, "3. Оттуда стрелка «маршрут к API» ведёт в greeting-service внутри k3s.")
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "Стрелка CI/CD · Helm сверху — выкладка кода разработчиком. "
        "Стрелка HTTPS снизу — живой трафик людей. Их специально развели.",
    )

    sty.add_heading2(doc, "3. Зачем он здесь вообще", "trf_s03_why")
    sty.add_normal(
        doc,
        "У Pod в Kubernetes нет нормального «сайта в интернете» из коробки. "
        "У него внутренний адрес кластера. С улицы его не видно.",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "Нужен кто-то с публичным IP, кто слушает обычные порты 80 и 443 и говорит:")
    sty.add_platform_block(
        doc,
        "Правило двери",
        "запрос на greeting-dev.example.com  →  отдай в greeting-service",
    )
    sty.add_normal(
        doc,
        "Это и есть задача Traefik. Без него пришлось бы либо открывать приложение "
        "голым NodePort на worker-ноде (неудобно, плохо с доменом и HTTPS), "
        "либо брать managed Ingress у облака — а в этом гайде облачный Ingress запрещён.",
    )
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "Аналогия: k3s — кухня ресторана. Traefik — швейцар у входной двери. "
        "Гость не идёт на кухню сам; швейцар встречает и провожает.",
    )

    sty.add_heading2(doc, "4. Почему на отдельном VPS", "trf_s04_vps")
    sty.add_normal(
        doc,
        "Три цветных блока на схеме — три разные машины с разными ролями.",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "Блок | Для кого | Что делает", consolas=True)
    sty.add_normal(doc, "Зелёный VPS devtools | разработчик | GitLab, сборка, Registry", consolas=True)
    sty.add_normal(doc, "Синий VPS k3s | приложение | крутит Pod с Java", consolas=True)
    sty.add_normal(doc, "Фиолетовый VPS Traefik | пользователь из интернета | принимает HTTPS и маршрутизирует", consolas=True)
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "GitLab не должен быть лицом сайта. Kubernetes не должен торчать в интернет "
        "всеми портами. Traefik — единственная точка, куда указывает DNS.",
    )
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "Подпись «отдельный кластер, не managed Ingress провайдера» значит: "
        "это не кнопка Ingress в панели Timeweb / AWS / Рег.облака. "
        "Это ваш собственный прокси на обычном VPS. Решение переносится на любого "
        "хостера: меняется только IP.",
    )

    sty.add_heading2(doc, "5. Почему именно Traefik", "trf_s05_why_traefik")
    sty.add_normal(
        doc,
        "Роль та же, что у NGINX Ingress в исходной Части 2: принять внешний HTTP(S) "
        "и довести до сервиса. Там это был Ingress внутри managed Kubernetes Timeweb. "
        "Здесь тот же вход вынесен на свои VPS.",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "Traefik для этого удобен, потому что заточен под маршрутизацию:")
    sty.add_normal(doc, "• Entrypoint — на каких портах слушать (:80, :443);")
    sty.add_normal(doc, "• Router — правило вроде «если Host = такой-то домен»;")
    sty.add_normal(doc, "• Service — куда дальше отправить (у вас — на NodePort worker-ноды k3s).")
    sty.add_citation(
        doc,
        "https://doc.traefik.io/traefik/",
        "Entrypoints are the network entry points into Traefik. They define the port that will "
        "receive the packets … Routers are in charge of connecting incoming requests to the "
        "services that can handle them. … Services are responsible for configuring how to reach "
        "the actual services that will eventually handle the incoming requests.",
        "Entrypoints — точки входа в сеть Traefik. Они задают порт, который принимает пакеты. "
        "Routers связывают входящие запросы с сервисами, которые могут их обработать. "
        "Services описывают, как достучаться до реальных сервисов, которые в итоге обработают запрос.",
    )
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "На том же месте можно было бы поставить nginx — роль двери не изменилась бы. "
        "Traefik выбран как отдельный, явно управляемый edge-прокси: домен, HTTPS, "
        "маршрут к API — всё в одном месте, без зависимости от Ingress-услуги облака.",
    )

    sty.add_heading2(doc, "6. Одна фраза вместо запутанной", "trf_s06_one_phrase")
    sty.add_normal(
        doc,
        "Вместо «внешний HTTP(S)-трафик принимает отдельный кластер Traefik» читайте так:",
    )
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "Пользователь из браузера никогда не ходит в Kubernetes напрямую. "
        "Он ходит на Traefik (публичные порты 80/443). Traefik смотрит имя сайта "
        "и пересылает запрос в greeting-service внутри k3s.",
    )
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "GitLab на этой картинке к пользовательскому трафику не относится. "
        "Он только принимает git push и через Helm выкатывает новую версию в синий блок.",
    )


def main() -> None:
    sty.reset_bookmarks()
    ensure_scheme_image()
    template = resolve_template()
    shutil.copy2(template, BUILD)
    doc = Document(BUILD)
    sty.clear_document_body(doc)
    build_document(doc)
    doc.save(BUILD)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(BUILD, OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
