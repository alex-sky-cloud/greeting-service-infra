# -*- coding: utf-8 -*-
"""
Генератор Word-гайда:
  greeting-service-guide - Часть 2M - CI-CD вручную (GitLab + k3s + Traefik).docx

Отличия от Части 2 (Timeweb + Terraform):
  - любой VPS-провайдер;
  - без Terraform и без managed Ingress облака;
  - только GitLab (Bitbucket не рассматривается);
  - Traefik — отдельный кластер на своих VPS;
  - S3 через MinIO;
  - всё через SSH / root / CLI.
"""
from __future__ import annotations

import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches

import docx_style as sty
from render_diagrams import (
    render_architecture,
    render_architecture_traffic,
    render_cicd,
    render_tech_map,
)

ROOT = Path(r"D:/Project_infra/greeting-service-infra")
HERE = Path(__file__).resolve().parent
OUTPUT = HERE / "greeting-service-guide - Часть 2M - CI-CD вручную (GitLab k3s Traefik).docx"
BUILD = HERE / "_build_chast2m.docx"
IMG_TECH = HERE / "images" / "tech-map.png"
IMG_ARCH = HERE / "images" / "architecture-manual.png"
IMG_TRAFFIC = HERE / "images" / "architecture-traffic.png"
IMG_CICD = HERE / "images" / "cicd-manual.png"
IMG_SCHEME11 = HERE / "images" / "scheme-1-1-chto-stroim-v2.jpg"

SCHEME11_CANDIDATES = [
    Path(
        r"C:/Users/sky/.cursor/projects/d-Project-infra-greeting-service-infra"
        r"/assets/c__Users_sky_AppData_Roaming_Cursor_User_workspaceStorage"
        r"_371f2630de332b8e39996771b99e1f6e_images_image-b38e4a38-d34b-49b8-b9e7-e84623b61d7e.jpg"
    ),
    HERE / "images" / "scheme-1-1-chto-stroim.png",
]

TEMPLATE_CANDIDATES = [
    ROOT / "docs/guide-deploy-terraform/Razdel-15a-gitlab-cicd.docx",
    ROOT / "docs/guide-deploy-terraform/Razdel-flyway-migrations.docx",
    ROOT / "docs/guide-deploy-terraform/Razdel-dns-ingress-dostup.docx",
]

# --- placeholders (пользователь подставляет свои IP) ---
DEVTOOLS_IP = "<DEVTOOLS_IP>"
K8S_MASTER_IP = "<K8S_MASTER_IP>"
K8S_WORKER_1_IP = "<K8S_WORKER_1_IP>"
TRAEFIK_1_IP = "<TRAEFIK_1_IP>"
TRAEFIK_2_IP = "<TRAEFIK_2_IP>"
TRAEFIK_FLOATING_IP = "<TRAEFIK_FLOATING_IP>"
STORAGE_IP = "<STORAGE_IP>"
POSTGRES_PRIMARY_IP = "<POSTGRES_PRIMARY_IP>"
POSTGRES_REPLICA_IP = "<POSTGRES_REPLICA_IP>"
DOMAIN = "greeting-dev.example.com"
GITLAB_GROUP = "greeting-group"
GITLAB_PROJECT = "greeting-service"
REPO_WIN = "/d/Project_infra/greeting-service-infra"
REPO_MAC = "~/Project_infra/greeting-service-infra"
REPO_LIN = "~/Project_infra/greeting-service-infra"

# Unique bookmarks across the WHOLE document (even if titles repeat)
TOC = [
    ("1. Архитектурное описание решения", "m2_s01_arch"),
    ("1.1. Что строим", "m2_s01_1_what"),
    ("1.1.1. Почему PostgreSQL не внутри k3s", "m2_s01_1_1_pg_outside"),
    ("1.2. Ключевые компоненты", "m2_s01_2_comp"),
    ("1.3. Ключевые архитектурные решения", "m2_s01_3_decisions"),
    ("1.4. Экскурсия по технологиям (зачем схема + текст)", "m2_s01_4_tour"),
    ("2. Схема архитектуры", "m2_s02_scheme"),
    ("2.1. Как читать схему", "m2_s02_1_howto"),
    ("2.2. Пояснение контуров на схеме", "m2_s02_2_contours"),
    ("2.3. Путь HTTP-запроса (отдельная схема)", "m2_s02_3_traffic"),
    ("2.4. Пояснение инструментов на схеме", "m2_s02_4_tools"),
    ("3. Схема CI/CD процесса", "m2_s03_cicd"),
    ("3.1. Как читать CI/CD-схему", "m2_s03_1_howto"),
    ("3.2. Пошаговый разбор pipeline", "m2_s03_2_steps"),
    ("4. Список серверов у провайдера", "m2_s04_servers"),
    ("5. Требования к локальному ПК", "m2_s05_local_pc"),
    ("5.1. Windows", "m2_s05_1_win"),
    ("5.2. macOS", "m2_s05_2_mac"),
    ("5.3. Linux (Ubuntu)", "m2_s05_3_linux"),
    ("6. SSH-доступ к серверам", "m2_s06_ssh"),
    ("6.1. Windows (Git Bash)", "m2_s06_1_win"),
    ("6.2. macOS (Terminal)", "m2_s06_2_mac"),
    ("6.3. Linux (Ubuntu)", "m2_s06_3_linux"),
    ("7. Подготовка серверов (root CLI)", "m2_s07_prepare"),
    ("8. Кластер Kubernetes (k3s)", "m2_s08_k3s"),
    ("8.1. Master", "m2_s08_1_master"),
    ("8.2. Workers", "m2_s08_2_workers"),
    ("8.3. kubeconfig на локальном ПК", "m2_s08_3_kubeconfig"),
    ("8.4. Разбор команд kubectl", "m2_s08_4_kubectl"),
    ("9. Кластер Traefik на отдельных VPS", "m2_s09_traefik"),
    ("9.1. Установка Traefik", "m2_s09_1_install"),
    ("9.2. Маршрут к приложению", "m2_s09_2_route"),
    ("10. Docker Registry на devtools", "m2_s10_registry"),
    ("11. GitLab CE на devtools", "m2_s11_gitlab"),
    ("12. GitLab Runner (self-hosted)", "m2_s12_runner"),
    ("13. S3 Bucket (MinIO)", "m2_s13_s3"),
    ("14. PostgreSQL на отдельных VPS (Patroni)", "m2_s14_pg"),
    ("15. Secrets и первый деплой Helm", "m2_s15_helm"),
    ("16. DNS", "m2_s16_dns"),
    ("17. GitLab CI/CD и подключение существующего репозитория", "m2_s17_pipeline"),
    ("17.1. Remote GitLab и push", "m2_s17_1_remote"),
    ("17.2. CI/CD Variables", "m2_s17_2_vars"),
    ("17.3. Пример ci/.gitlab-ci.yml", "m2_s17_3_yml"),
    ("18. Проверка результата", "m2_s18_verify"),
    ("19. Типичные ошибки", "m2_s19_errors"),
    ("20. Финальная сводка", "m2_s20_summary"),
]


def resolve_template() -> Path:
    for p in TEMPLATE_CANDIDATES:
        if p.exists():
            return p
    raise FileNotFoundError("Не найден Word-шаблон со стилями Heading")


def add_toc(doc: Document) -> None:
    sty.add_heading2(doc, "Оглавление", "m2_toc")
    for title, anchor in TOC:
        sty.add_hyperlink_paragraph(doc, title, anchor, indent=sty.is_toc_subsection(title))
    sty.add_empty(doc)


def add_figure(doc: Document, path: Path, caption: str) -> None:
    sty.add_empty(doc)
    sty.add_normal(doc, caption, consolas=True)
    sty.add_empty(doc)
    # Шире + крупные PNG-шрифты: текст на странице читается без лупы
    doc.add_picture(str(path), width=Inches(6.7))
    sty.add_empty(doc)


def add_os_triple(doc: Document, win: str, mac: str, linux: str) -> None:
    sty.add_platform_block(doc, "Локальный ПК — Windows (Git Bash)", win)
    sty.add_platform_block(doc, "Локальный ПК — macOS (Terminal)", mac)
    sty.add_platform_block(doc, "Локальный ПК — Linux / Ubuntu (Terminal)", linux)


def ensure_scheme11() -> Path:
    if IMG_SCHEME11.exists():
        return IMG_SCHEME11
    IMG_SCHEME11.parent.mkdir(parents=True, exist_ok=True)
    for src in SCHEME11_CANDIDATES:
        if src.exists():
            shutil.copy2(src, IMG_SCHEME11)
            return IMG_SCHEME11
    return IMG_SCHEME11


def build_document(doc: Document) -> None:
    sty.add_normal(doc, "ЧАСТЬ II-M", consolas=True)
    sty.add_normal(doc, "CI/CD и развёртывание вручную", consolas=True)
    sty.add_normal(
        doc,
        "GitLab CE · Runner · Docker Registry · k3s · Traefik edge · MinIO S3 · любой VPS-провайдер",
        consolas=True,
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "Java-микросервис: ручная инфраструктура без Terraform")
    sty.add_normal(doc, "Версия: 1.0 | 2026-08 | Целевая аудитория: backend developer middle+")
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "Все описания — на русском языке. Технические термины — на английском. Команды — на английском.",
    )
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "Отличие от Части 2 (Timeweb + Terraform): инфраструктура создаётся только через панель "
        "провайдера (заказ VPS) и root/SSH/CLI. Terraform, cloud-plugins и Bitbucket не используются. "
        "Вход из интернета — через отдельный кластер Traefik на своих VPS, а не через managed Ingress облака.",
    )
    sty.add_empty(doc)
    add_toc(doc)

    # ── 1 ────────────────────────────────────────────────────────────────
    sty.add_heading2(doc, "1. Архитектурное описание решения", "m2_s01_arch")

    sty.add_heading3(doc, "1.1. Что строим", "m2_s01_1_what")
    sty.add_normal(
        doc,
        "Ниже — описание той же схемы 1.1: отдельные VPS для CI/CD, приложения (k3s), "
        "входа из интернета (Traefik) и базы данных (PostgreSQL). "
        "«Кластер базы» здесь — это primary + replica на своих машинах, а не база внутри Kubernetes.",
    )
    if IMG_SCHEME11.exists():
        add_figure(doc, IMG_SCHEME11, "Рисунок 1.1. Что строим: k3s только для приложения, PostgreSQL на отдельных VPS")
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "1. Простой Java-микросервис на Spring Boot с REST endpoint GET /api/greeting "
        "(репозиторий greeting-service-infra у вас уже есть).",
    )
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "2. Сервис разворачивается в самостоятельно установленном Kubernetes-кластере (k3s) "
        "на обычных VPS любого облачного/VPS-провайдера. На схеме блок k3s — это worker-ноды, "
        "где крутятся Pod приложения. Доступ из интернета идёт не напрямую в k3s, а через Traefik.",
    )
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "3. Весь путь от git push до новой версии в кластере автоматизирован через GitLab CI/CD "
        "(self-hosted GitLab CE + self-hosted GitLab Runner) и Helm.",
    )
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "4. GitLab CE, Docker Registry и GitLab Runner размещаются на отдельном VPS (devtools): "
        "сборка, тесты, публикация образа.",
    )
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "5. Внешний HTTP(S)-трафик принимает отдельный кластер Traefik на своих VPS "
        "(порты :80 / :443). Это не managed Ingress провайдера. Traefik маршрутизирует запрос "
        "к API greeting-service внутри k3s.",
    )
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "6. База данных — отдельный PostgreSQL-кластер на своих VPS (не внутри k3s): "
        "один primary, одна или несколько replica. Приложение на worker-нодах ходит в базу "
        "по частной сети. Диск каждой машины принадлежит только этой ноде базы.",
    )

    sty.add_heading3(doc, "1.1.1. Почему PostgreSQL не внутри k3s", "m2_s01_1_1_pg_outside")
    sty.add_normal(
        doc,
        "Kubernetes умеет останавливать контейнеры и запускать их на другом узле. "
        "Для сайта (stateless Pod) это нормально: потерялся экземпляр — поднялся новый.",
    )
    sty.add_citation(
        doc,
        "https://kubernetes.io/docs/concepts/workloads/pods/",
        "You'll rarely create individual Pods directly in Kubernetes—even singleton Pods. "
        "This is because Pods are designed as relatively ephemeral, disposable entities. "
        "When a Pod gets created … The Pod remains on that node until the Pod finishes execution, "
        "the Pod object is deleted, the Pod is evicted for lack of resources, or the node fails.",
        "Поды в Kubernetes редко создают напрямую даже по одному. Поды задуманы как относительно "
        "эфемерные, одноразовые сущности. Под остаётся на узле, пока не завершится, не будет удалён, "
        "не будет вытеснен из‑за нехватки ресурсов или пока не упадёт узел.",
    )
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "Для PostgreSQL этого недостаточно. У базы данные на диске, открытые клиентские соединения "
        "и журнал записей (WAL). WAL — это не «лог для красоты»: изменения в файлах таблиц "
        "записываются только после того, как описание этих изменений сброшено на постоянное хранилище. "
        "Если диск «уехал» не туда или одновременно живы два primary, можно получить порчу данных, "
        "а не «магический отказоустойчивый кластер».",
    )
    sty.add_citation(
        doc,
        "https://www.postgresql.org/docs/current/wal-intro.html",
        "WAL's central concept is that changes to data files (where tables and indexes reside) "
        "must be written only after those changes have been logged, that is, after WAL records "
        "describing the changes have been flushed to permanent storage.",
        "Центральная идея WAL: изменения в файлах данных (таблицы и индексы) должны записываться "
        "только после того, как эти изменения занесены в журнал, то есть после сброса записей WAL "
        "на постоянное хранилище.",
    )
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "Поэтому кластер базы делают иначе, чем кластер приложения. Два-три отдельных VPS, "
        "у каждого свой диск. Один сервер — primary (чтение и запись), остальные — replica "
        "(следят за primary). Если одна машина вышла из строя, живые остаются; replica можно "
        "повысить до primary. Это и есть high availability в терминах PostgreSQL.",
    )
    sty.add_citation(
        doc,
        "https://www.postgresql.org/docs/current/high-availability.html",
        "Database servers can work together to allow a second server to take over quickly "
        "if the primary server fails (high availability) … Servers that can modify data are "
        "called read/write, master or primary servers. Servers that track changes in the primary "
        "are called standby or secondary servers.",
        "Серверы баз данных могут работать вместе, чтобы второй сервер быстро принял нагрузку, "
        "если primary отказал (высокая доступность). Серверы, которые могут изменять данные, "
        "называют primary. Серверы, которые отслеживают изменения primary, называют standby / replica.",
    )
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "Два контейнера Postgres внутри k3s на тех же worker-нодах, что и приложение, — это не "
        "боевой кластер базы. Упала одна worker-машина или общее хранилище — судьба обеих копий "
        "общая. Плюс диск в Kubernetes часто сетевой или разделяемый: скачки задержки для WAL "
        "сразу бьют по Postgres. Соседние сервисы на том же узле дают noisy neighbor: нагрузка "
        "на диск и CPU сайта не должна тормозить записи в базу. Обновление k3s не должно быть "
        "окном обслуживания самой БД.",
    )
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "Ещё одно уточнение, чтобы не путать слова. Kubernetes даёт перезапуск пода: процесс "
        "подняли заново. Failover базы — это другое: старый primary нужно изолировать (fencing), "
        "replica — повысить, клиентам указать новый адрес записи. Перезапуск контейнера сам по себе "
        "этого не делает. Для такого контура на отдельных машинах обычно берут Patroni: он "
        "управляет PostgreSQL и автоматическим переключением primary.",
    )
    sty.add_citation(
        doc,
        "https://patroni.readthedocs.io/en/latest/",
        "Patroni is a template for high availability (HA) PostgreSQL solutions using Python. "
        "For maximum accessibility, Patroni supports a variety of distributed configuration stores "
        "like ZooKeeper, etcd, Consul or Kubernetes.",
        "Patroni — шаблон решений высокой доступности PostgreSQL на Python. Для гибкости Patroni "
        "поддерживает разные распределённые хранилища конфигурации: ZooKeeper, etcd, Consul или Kubernetes.",
    )
    sty.add_citation(
        doc,
        "https://patroni.readthedocs.io/en/latest/faq.html",
        "If a primary node fails, Patroni will not only fail over to a replica, but also attempt "
        "to rejoin the former primary as a replica of the new primary. … You should not attempt "
        "to manage Postgres directly! Any attempt of bouncing the Postgres server without Patroni "
        "can lead your cluster to face failovers.",
        "Если primary отказывает, Patroni не только делает failover на replica, но и пытается "
        "вернуть бывший primary уже как replica нового primary. Postgres нельзя дергать напрямую: "
        "перезапуск сервера в обход Patroni может привести к лишним failover.",
    )
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "Замечание по терминологии: Patroni умеет работать и поверх Kubernetes (это прямо сказано "
        "в его документации). В этом гайде мы сознательно этого не делаем: база живёт на отдельных "
        "VPS со своими дисками, k3s остаётся контуром приложения. Связь — частная сеть, как на схеме.",
    )

    sty.add_heading3(doc, "1.2. Ключевые компоненты", "m2_s01_2_comp")
    sty.add_normal(doc, "• Исходный код — Java 21, Spring Boot; хранится в GitLab CE на VPS devtools;", consolas=True)
    sty.add_normal(doc, "• CI/CD — GitLab Pipelines + self-hosted Runner (shell executor);", consolas=True)
    sty.add_normal(doc, "• Docker Registry — distribution/registry:2 на devtools :5000;", consolas=True)
    sty.add_normal(doc, "• Kubernetes — k3s: worker-ноды для приложения (Pod greeting-service);", consolas=True)
    sty.add_normal(doc, "• Вход из интернета — Traefik Proxy на отдельном edge-кластере VPS;", consolas=True)
    sty.add_normal(doc, "• БД — PostgreSQL primary + replica на отдельных VPS, Patroni, частная сеть;", consolas=True)
    sty.add_normal(doc, "• S3 — MinIO на VPS storage (совместимый S3 API);", consolas=True)
    sty.add_normal(doc, "• Деплой — Helm 3 из Runner / с локального ПК;", consolas=True)
    sty.add_normal(doc, "• IaC — нет; всё вручную через CLI и root.", consolas=True)

    sty.add_heading3(doc, "1.3. Ключевые архитектурные решения", "m2_s01_3_decisions")
    sty.add_normal(doc, "Почему GitLab CE на VPS, а не GitLab.com SaaS?")
    sty.add_normal(
        doc,
        "Задача требует автономной инфраструктуры: репозиторий, Runner и Registry "
        "полностью под вашим контролем на своих серверах.",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "Почему Traefik на отдельном VPS-кластере, а не Ingress внутри k3s от облака?")
    sty.add_normal(
        doc,
        "Managed Ingress облака привязывает вас к провайдеру. Отдельный Traefik-кластер "
        "переносится вместе с конфигурацией на любого хостера: меняются только IP серверов.",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "Документация Traefik (установка):")
    sty.add_normal(doc, "https://doc.traefik.io/traefik/getting-started/install-traefik/")
    sty.add_citation(
        doc,
        "https://doc.traefik.io/traefik/getting-started/install-traefik/",
        "Traefik can be deployed in various environments. Choose your preferred deployment method: "
        "Kubernetes Quick Start — Deploy Traefik using Helm; Docker Quick Start — Deploy Traefik using Docker.",
        "Traefik можно развернуть в разных окружениях. Выберите способ: быстрый старт в Kubernetes "
        "(Helm) или быстрый старт в Docker.",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "Почему k3s, а не managed Kubernetes?")
    sty.add_normal(
        doc,
        "k3s ставится одним скриптом на Ubuntu и даёт полноценный Kubernetes API без зависимости "
        "от панели конкретного облака.",
    )
    sty.add_citation(
        doc,
        "https://docs.k3s.io/quick-start",
        "K3s provides an installation script that is a convenient way to install it as a service "
        "on systemd or openrc based systems. This script is available at https://get.k3s.io. "
        "To install K3s using this method, just run: curl -sfL https://get.k3s.io | sh -",
        "K3s предоставляет установочный скрипт для удобной установки как сервиса на системах "
        "с systemd или openrc. Скрипт доступен на https://get.k3s.io. Установка: "
        "curl -sfL https://get.k3s.io | sh -",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "Почему MinIO для S3?")
    sty.add_normal(
        doc,
        "MinIO реализует S3-совместимый API на своём VPS: bucket, AccessKey и SecretKey "
        "настраиваются вручную, без Terraform-ресурса twc_s3_bucket.",
    )
    sty.add_citation(
        doc,
        "https://min.io/docs/minio/linux/index.html",
        "MinIO is a high-performance, S3 compatible object store.",
        "MinIO — высокопроизводительное объектное хранилище, совместимое с S3.",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "Почему PostgreSQL на отдельных VPS, а не StatefulSet в k3s?")
    sty.add_normal(
        doc,
        "См. п. 1.1.1. Коротко: сайт можно переносить между узлами; у базы диск, WAL и роль primary. "
        "Два Postgres-контейнера на worker-нодах k3s не дают настоящего failover. "
        "Для боевого контура — отдельные машины и Patroni.",
    )

    sty.add_heading3(doc, "1.4. Экскурсия по технологиям (зачем схема + текст)", "m2_s01_4_tour")
    sty.add_normal(
        doc,
        "Вы правы в подходе: технологию проще понять, когда рядом есть функциональная схема "
        "и короткое пояснение «зачем этот блок», а не только сухой список команд. "
        "Ниже — экскурсия-карта. Дальше в разделах 2 и 3 каждая схема снова сопровождается разбором.",
    )
    add_figure(doc, IMG_TECH, "Рисунок 1.4. Экскурсия: роль каждой технологии (крупный шрифт)")
    sty.add_normal(doc, "Как пользоваться этой картой:")
    sty.add_normal(doc, "1. Сначала прочитайте аналогию на схеме (сейф, завод, склад, швейцар).")
    sty.add_normal(doc, "2. Потом откройте схему архитектуры (раздел 2) — те же блоки уже в «боевой» раскладке.")
    sty.add_normal(doc, "3. Команды CLI в разделах 7–17 — это уже «как настроить руками», после того как понятна роль.")

    # ── 2 ────────────────────────────────────────────────────────────────
    sty.add_heading2(doc, "2. Схема архитектуры", "m2_s02_scheme")
    sty.add_normal(
        doc,
        "Схема нарисована крупными блоками сверху вниз (те же цвета, что в Части 2: "
        "тёплый локальный ПК/devtools, бирюзовый Kubernetes, сиреневый edge/интернет). "
        "Mermaid не используется.",
    )
    add_figure(doc, IMG_ARCH, "Рисунок 2.1. Архитектура — крупный обзор (читайте сверху вниз)")

    sty.add_heading3(doc, "2.1. Как читать схему", "m2_s02_1_howto")
    sty.add_normal(doc, "1. Идите по номерам 1→5: от вашего ПК к интернету и хранилищу.")
    sty.add_empty(doc)
    sty.add_normal(doc, "2. Внизу схемы два потока — запомните их как каркас всего гайда:")
    sty.add_normal(doc, "   A) поток кода (CI/CD);", consolas=True)
    sty.add_normal(doc, "   B) поток трафика (запрос пользователя).", consolas=True)
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "3. Если блок на схеме непонятен — вернитесь к рисунку 1.4 (экскурсия) и снова к тексту ниже.",
    )

    sty.add_heading3(doc, "2.2. Пояснение контуров на схеме", "m2_s02_2_contours")
    sty.add_normal(doc, "Локальный ПК разработчика", consolas=True)
    sty.add_normal(
        doc,
        "Здесь вы пишете код и отправляете его в Git. kubectl/helm нужны, чтобы проверить кластер "
        "или выкатить вручную. Terraform на схеме намеренно отсутствует: серверы заказываете в панели, "
        "настраиваете по SSH.",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "VPS «devtools» — фабрика CI/CD", consolas=True)
    sty.add_normal(
        doc,
        "GitLab CE принимает git push и хранит историю. Runner — это процесс на том же сервере, "
        "который выполняет jobs из pipeline. Docker Registry хранит собранные образы, чтобы "
        "Kubernetes мог их скачать (image pull).",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "VPS-кластер Kubernetes (k3s)", consolas=True)
    sty.add_normal(
        doc,
        "k3s — это полноценный Kubernetes в компактной упаковке для обычных VPS. "
        "На схеме 1.1 блок k3s — worker-ноды, где крутится greeting-service. "
        "Service NodePort — «дверь», в которую стучится Traefik. "
        "Базы данных в этом блоке нет.",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "VPS PostgreSQL (отдельный кластер)", consolas=True)
    sty.add_normal(
        doc,
        "Жёлтый блок схемы: primary и replica на своих VPS, каждый со своим диском. "
        "Связь с приложением — частная сеть, не публичный интернет.",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "Отдельный кластер Traefik (edge)", consolas=True)
    sty.add_normal(
        doc,
        "Это вход из интернета. Не путайте с Ingress-кнопкой в панели облака: Traefik стоит на "
        "своих двух VPS, конфиг лежит у вас в файлах. Floating IP даёт стабильный адрес для DNS.",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "VPS «storage» / MinIO", consolas=True)
    sty.add_normal(
        doc,
        "S3-совместимое хранилище без Terraform-ресурса облака. Bucket и ключи настраиваете CLI-утилитой mc.",
    )

    sty.add_heading3(doc, "2.3. Путь HTTP-запроса (отдельная схема)", "m2_s02_3_traffic")
    sty.add_normal(
        doc,
        "Архитектурный обзор показывает «где что стоит». Ниже — отдельная функциональная схема "
        "только для одного пользовательского запроса. Так легче понять Traefik.",
    )
    add_figure(doc, IMG_TRAFFIC, "Рисунок 2.3. Путь GET /api/greeting (крупно, слева направо)")
    sty.add_normal(doc, "Разбор по шагам схемы 1б:")
    sty.add_normal(doc, "1. Браузер спрашивает DNS: «какой IP у greeting-dev.example.com?»")
    sty.add_normal(doc, "2. DNS отвечает Floating IP Traefik.")
    sty.add_normal(doc, "3. Traefik читает заголовок Host и выбирает backend.")
    sty.add_normal(doc, "4. Backend = IP worker-ноды k3s + NodePort сервиса.")
    sty.add_normal(doc, "5. Запрос попадает в Pod Spring Boot → ответ JSON/текст greeting.")
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "Аналогия со схемы: Traefik — швейцар. Kubernetes — этажи с номерами (Pod). "
        "Без швейцара гость не знает, в какую дверь стучать с улицы.",
    )

    sty.add_heading3(doc, "2.4. Пояснение инструментов на схеме", "m2_s02_4_tools")
    sty.add_normal(doc, "Что делает GitLab CE")
    sty.add_normal(
        doc,
        "Хранит git-репозиторий и описывает автоматизацию в ci/.gitlab-ci.yml. "
        "После push создаёт pipeline и отдаёт jobs Runner’у.",
    )
    sty.add_citation(
        doc,
        "https://docs.gitlab.com/ci/pipelines/",
        "A pipeline is a top-level component for continuous integration, delivery, and deployment.",
        "Pipeline — верхнеуровневый компонент непрерывной интеграции, доставки и развёртывания.",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "Что делает GitLab Runner")
    sty.add_normal(
        doc,
        "Исполнитель jobs. В нашем гайде executor = shell: команды идут прямо на devtools "
        "(нужны docker, JDK, helm, kubectl на этом хосте).",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "Что делает Docker Registry")
    sty.add_normal(
        doc,
        "Принимает docker push и отдаёт docker pull / image pull из Kubernetes. "
        "Без registry кластер не узнает, откуда взять новый образ.",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "Что делает Helm")
    sty.add_normal(
        doc,
        "Пакетный менеджер Kubernetes: chart + values → Deployment/Service/и т.д. "
        "Стрелка helm upgrade на схемах = «выкати или обнови приложение».",
    )
    sty.add_citation(
        doc,
        "https://helm.sh/docs/",
        "Helm helps you manage Kubernetes applications — Helm Charts help you define, install, and upgrade even the most complex Kubernetes application.",
        "Helm помогает управлять приложениями Kubernetes — Helm Charts помогают описывать, устанавливать и обновлять даже сложные приложения.",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "Что делает kubectl")
    sty.add_normal(
        doc,
        "CLI для диагностики и ручного управления: поды, логи, describe, apply. "
        "Если Helm «ставит», то kubectl помогает понять, что пошло не так.",
    )
    sty.add_citation(
        doc,
        "https://kubernetes.io/docs/reference/kubectl/",
        "kubectl controls the Kubernetes cluster manager.",
        "kubectl управляет менеджером кластера Kubernetes.",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "Что делает Traefik")
    sty.add_normal(
        doc,
        "Обратный прокси / edge-маршрутизатор. В этом гайде он вынесен на отдельные VPS, "
        "чтобы не зависеть от managed Ingress провайдера.",
    )
    sty.add_citation(
        doc,
        "https://doc.traefik.io/traefik/",
        "Here’s how it works—Traefik receives requests on behalf of your system, identifies which "
        "components are responsible for handling them, and routes them securely.",
        "Как это работает: Traefik принимает запросы от имени вашей системы, определяет, "
        "какие компоненты должны их обработать, и безопасно маршрутизирует их.",
    )
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "Важно: Terraform на схемах отсутствует намеренно. Платформу (VPS, IP, ОС) готовите вручную; "
        "Helm и kubectl работают уже поверх готового Kubernetes.",
    )

    # ── 3 ────────────────────────────────────────────────────────────────
    sty.add_heading2(doc, "3. Схема CI/CD процесса", "m2_s03_cicd")
    sty.add_normal(
        doc,
        "Ниже — вертикальная схема pipeline. Каждый шаг крупным шрифтом; рядом в тексте — "
        "что происходит «под капотом» и зачем этот шаг нужен.",
    )
    add_figure(doc, IMG_CICD, "Рисунок 3.1. CI/CD flow (только GitLab) — шаги сверху вниз")

    sty.add_heading3(doc, "3.1. Как читать CI/CD-схему", "m2_s03_1_howto")
    sty.add_normal(doc, "1. Идите сверху вниз: от git push к curl-проверке.")
    sty.add_normal(doc, "2. Bitbucket на схеме нет — только GitLab CE + self-hosted Runner.")
    sty.add_normal(doc, "3. Traefik на этой схеме в конце: он не собирает код, он пускает пользователей к уже выкатанному Pod.")

    sty.add_heading3(doc, "3.2. Пошаговый разбор pipeline", "m2_s03_2_steps")
    sty.add_normal(doc, "Шаг 1. Developer → git push", consolas=True)
    sty.add_normal(doc, "Вы фиксируете изменения локально и отправляете ветку develop на GitLab.")
    sty.add_empty(doc)
    sty.add_normal(doc, "Шаг 2. GitLab CE стартует pipeline", consolas=True)
    sty.add_normal(
        doc,
        "GitLab читает ci/.gitlab-ci.yml и ставит jobs в очередь Runner’у с тегом self-hosted.",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "Шаг 3. Runner: test + docker", consolas=True)
    sty.add_normal(
        doc,
        "./gradlew test проверяет код. docker build упаковывает приложение. "
        "docker push кладёт образ в Registry с тегом short SHA коммита.",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "Шаг 4. Registry хранит версию", consolas=True)
    sty.add_normal(doc, "Каждый успешный pipeline оставляет воспроизводимый артефакт-образ.")
    sty.add_empty(doc)
    sty.add_normal(doc, "Шаг 5. Helm обновляет k3s", consolas=True)
    sty.add_normal(
        doc,
        "helm upgrade --install меняет Deployment: Kubernetes скачивает новый образ и "
        "поднимает Pod. readinessProbe не пускает трафик, пока приложение не готово.",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "Шаг 6. Проверка снаружи через Traefik", consolas=True)
    sty.add_normal(
        doc,
        "Если DNS и Host-правило уже настроены, curl по доменному имени показывает новую версию. "
        "Если DNS ещё нет — проверяйте через Floating IP и заголовок Host (раздел 16).",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "Правила веток (как на схеме):")
    sty.add_normal(doc, "• develop — авто-деплой в namespace dev;")
    sty.add_normal(doc, "• main — deploy-prod только вручную в UI GitLab;")
    sty.add_normal(doc, "• feature/* — только build/test.")
    sty.add_citation(
        doc,
        "https://docs.gitlab.com/ci/pipelines/",
        "A pipeline is a top-level component for continuous integration, delivery, and deployment.",
        "Pipeline — верхнеуровневый компонент непрерывной интеграции, доставки и развёртывания.",
    )

    # ── 4 ────────────────────────────────────────────────────────────────
    sty.add_heading2(doc, "4. Список серверов у провайдера", "m2_s04_servers")
    sty.add_normal(
        doc,
        "У любого VPS-провайдера заказываете только виртуальные серверы Ubuntu 22.04 LTS "
        "и (желательно) один Floating IP для Traefik. Managed Kubernetes / managed Ingress / Terraform — не нужны.",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "Сервер | vCPU | RAM | Диск | Роль", consolas=True)
    sty.add_normal(doc, "devtools | 4 | 8 ГБ | 100 ГБ | GitLab CE + Registry + Runner", consolas=True)
    sty.add_normal(doc, "k8s-master | 2 | 4 ГБ | 50 ГБ | control-plane k3s", consolas=True)
    sty.add_normal(doc, "k8s-worker-1 | 2 | 4 ГБ | 50 ГБ | worker k3s (приложение)", consolas=True)
    sty.add_normal(doc, "k8s-worker-2 | 2 | 4 ГБ | 50 ГБ | worker k3s (приложение)", consolas=True)
    sty.add_normal(doc, "traefik-1 | 2 | 2 ГБ | 40 ГБ | Traefik node 1", consolas=True)
    sty.add_normal(doc, "traefik-2 | 2 | 2 ГБ | 40 ГБ | Traefik node 2 (HA)", consolas=True)
    sty.add_normal(doc, "postgres-1 | 2 | 4 ГБ | 80 ГБ | PostgreSQL primary", consolas=True)
    sty.add_normal(doc, "postgres-2 | 2 | 4 ГБ | 80 ГБ | PostgreSQL replica", consolas=True)
    sty.add_normal(doc, "storage | 2 | 4 ГБ | 100 ГБ | MinIO (S3)", consolas=True)
    sty.add_empty(doc)
    sty.add_normal(doc, "Алгоритм заказа (одинаков у большинства панелей):")
    sty.add_normal(doc, "1. Войти в панель провайдера → Создать сервер / New instance.")
    sty.add_normal(doc, "2. ОС: Ubuntu 22.04 LTS.")
    sty.add_normal(doc, "3. Добавить публичный SSH-ключ.")
    sty.add_normal(doc, "4. Заказать Floating IP и привязать к traefik-1 (или держать на keepalived VIP).")
    sty.add_normal(doc, "5. Записать все публичные IP в файл infra-servers.env (см. ниже).")
    sty.add_platform_block(
        doc,
        "Локальный ПК — любой ОС, файл переменных",
        "cat > infra-servers.env << 'EOF'\n"
        f"DEVTOOLS_IP={DEVTOOLS_IP}\n"
        f"K8S_MASTER_IP={K8S_MASTER_IP}\n"
        f"K8S_WORKER_1_IP={K8S_WORKER_1_IP}\n"
        f"K8S_WORKER_2_IP=<K8S_WORKER_2_IP>\n"
        f"TRAEFIK_1_IP={TRAEFIK_1_IP}\n"
        f"TRAEFIK_2_IP={TRAEFIK_2_IP}\n"
        f"TRAEFIK_FLOATING_IP={TRAEFIK_FLOATING_IP}\n"
        f"POSTGRES_PRIMARY_IP={POSTGRES_PRIMARY_IP}\n"
        f"POSTGRES_REPLICA_IP={POSTGRES_REPLICA_IP}\n"
        f"STORAGE_IP={STORAGE_IP}\n"
        "EOF",
    )

    # ── 5 ────────────────────────────────────────────────────────────────
    sty.add_heading2(doc, "5. Требования к локальному ПК", "m2_s05_local_pc")
    sty.add_normal(doc, "Нужны: git, ssh-клиент, kubectl, helm, (опционально) Docker Desktop.")

    sty.add_heading3(doc, "5.1. Windows", "m2_s05_1_win")
    sty.add_normal(
        doc,
        "Рекомендуется Git for Windows + Git Bash: пути и команды ближе к Linux/macOS. "
        "kubectl/helm — бинарники в PATH или через winget/chocolatey.",
    )
    sty.add_platform_block(
        doc,
        "Локальный ПК — Windows (PowerShell / cmd)",
        "winget install Git.Git\n"
        "winget install Kubernetes.kubectl\n"
        "# Helm: скачайте release с https://github.com/helm/helm/releases и добавьте helm.exe в PATH",
    )

    sty.add_heading3(doc, "5.2. macOS", "m2_s05_2_mac")
    sty.add_platform_block(
        doc,
        "Локальный ПК — macOS (Terminal)",
        "brew install git kubectl helm\n"
        "ssh -V",
    )

    sty.add_heading3(doc, "5.3. Linux (Ubuntu)", "m2_s05_3_linux")
    sty.add_platform_block(
        doc,
        "Локальный ПК — Linux / Ubuntu",
        "sudo apt-get update\n"
        "sudo apt-get install -y git curl openssh-client\n"
        "curl -LO \"https://dl.k8s.io/release/$(curl -L -s https://dl.k8s.io/release/stable.txt)/bin/linux/amd64/kubectl\"\n"
        "chmod +x kubectl && sudo mv kubectl /usr/local/bin/\n"
        "curl https://raw.githubusercontent.com/helm/helm/main/scripts/get-helm-3 | bash",
    )
    sty.add_citation(
        doc,
        "https://kubernetes.io/docs/tasks/tools/",
        "kubectl is a command line tool that lets you control Kubernetes clusters.",
        "kubectl — утилита командной строки для управления кластерами Kubernetes.",
    )

    # ── 6 ────────────────────────────────────────────────────────────────
    sty.add_heading2(doc, "6. SSH-доступ к серверам", "m2_s06_ssh")
    sty.add_normal(
        doc,
        "На серверах пользователь root (или ubuntu с sudo — тогда добавьте sudo). "
        "Ключ: ~/.ssh/id_ed25519. Ниже — три официальных варианта подключения.",
    )

    sty.add_heading3(doc, "6.1. Windows (Git Bash)", "m2_s06_1_win")
    sty.add_platform_block(
        doc,
        "Локальный ПК — Windows (Git Bash)",
        "source ./infra-servers.env\n"
        "ssh -i /c/Users/$USER/.ssh/id_ed25519 root@${DEVTOOLS_IP} \"echo connected\"\n"
        "ssh -i /c/Users/$USER/.ssh/id_ed25519 root@${K8S_MASTER_IP} \"echo connected\"\n"
        "ssh -i /c/Users/$USER/.ssh/id_ed25519 root@${TRAEFIK_1_IP} \"echo connected\"",
    )
    sty.add_normal(
        doc,
        "Пояснение: -i задаёт приватный ключ; root@IP — пользователь и хост; "
        "кавычки запускают удалённую команду без интерактивной сессии.",
    )

    sty.add_heading3(doc, "6.2. macOS (Terminal)", "m2_s06_2_mac")
    sty.add_platform_block(
        doc,
        "Локальный ПК — macOS (Terminal)",
        "source ./infra-servers.env\n"
        "ssh -i ~/.ssh/id_ed25519 root@${DEVTOOLS_IP} \"echo connected\"\n"
        "ssh -i ~/.ssh/id_ed25519 root@${K8S_MASTER_IP} \"echo connected\"\n"
        "ssh -i ~/.ssh/id_ed25519 root@${TRAEFIK_1_IP} \"echo connected\"",
    )

    sty.add_heading3(doc, "6.3. Linux (Ubuntu)", "m2_s06_3_linux")
    sty.add_platform_block(
        doc,
        "Локальный ПК — Linux / Ubuntu",
        "source ./infra-servers.env\n"
        "ssh -i ~/.ssh/id_ed25519 root@${DEVTOOLS_IP} \"echo connected\"\n"
        "ssh -i ~/.ssh/id_ed25519 root@${K8S_MASTER_IP} \"echo connected\"\n"
        "ssh -i ~/.ssh/id_ed25519 root@${TRAEFIK_1_IP} \"echo connected\"",
    )
    sty.add_normal(doc, "Успех: строка connected без запроса пароля.")

    # ── 7 ────────────────────────────────────────────────────────────────
    sty.add_heading2(doc, "7. Подготовка серверов (root CLI)", "m2_s07_prepare")
    sty.add_normal(doc, "На каждом сервере после первого SSH:")
    sty.add_platform_block(
        doc,
        "На сервере (после SSH), Ubuntu root",
        "apt-get update\n"
        "apt-get upgrade -y\n"
        "apt-get install -y curl wget gnupg git ufw ca-certificates\n"
        "hostnamectl set-hostname <ROLE_NAME>   # уникальное имя: k8s-master, traefik-1, ...\n"
        "ufw allow OpenSSH\n"
        "ufw allow 80/tcp\n"
        "ufw allow 443/tcp\n"
        "ufw --force enable",
    )
    sty.add_normal(doc, "На devtools, traefik-*, storage дополнительно Docker:")
    sty.add_platform_block(
        doc,
        "На сервере (devtools / traefik / storage)",
        "curl -fsSL https://get.docker.com | sh\n"
        "systemctl enable --now docker\n"
        "docker --version",
    )
    sty.add_normal(
        doc,
        "Ключи: apt-get update — обновить индексы пакетов; install -y — без вопросов; "
        "ufw allow — открыть порт; systemctl enable --now — автозапуск и старт сервиса.",
    )

    # ── 8 ────────────────────────────────────────────────────────────────
    sty.add_heading2(doc, "8. Кластер Kubernetes (k3s)", "m2_s08_k3s")
    sty.add_normal(
        doc,
        "k3s ставится официальным скриптом get.k3s.io. Traefik, встроенный в k3s по умолчанию, "
        "отключаем: внешний вход будет через отдельный Traefik-кластер.",
    )

    sty.add_heading3(doc, "8.1. Master", "m2_s08_1_master")
    sty.add_platform_block(
        doc,
        "На k8s-master (после SSH)",
        "curl -sfL https://get.k3s.io | INSTALL_K3S_EXEC=\"--disable traefik\" sh -\n"
        "k3s kubectl get nodes\n"
        "cat /var/lib/rancher/k3s/server/node-token",
    )
    sty.add_normal(
        doc,
        "Пояснение флагов: curl -sfL — silent/fail/follow redirects; INSTALL_K3S_EXEC="
        "\"--disable traefik\" — не ставить встроенный Traefik k3s; sh - — выполнить скрипт.",
    )
    sty.add_citation(
        doc,
        "https://docs.k3s.io/installation/network-options",
        "If you wish to use an alternate Ingress controller, start K3s with --disable=traefik.",
        "Если нужен другой Ingress-контроллер, запустите K3s с --disable=traefik.",
    )

    sty.add_heading3(doc, "8.2. Workers", "m2_s08_2_workers")
    sty.add_platform_block(
        doc,
        "На k8s-worker-1 / k8s-worker-2 (после SSH)",
        f"curl -sfL https://get.k3s.io | K3S_URL=https://{K8S_MASTER_IP}:6443 K3S_TOKEN=<NODE_TOKEN> sh -\n"
        "# на master проверьте:\n"
        "k3s kubectl get nodes -o wide",
    )
    sty.add_normal(
        doc,
        "K3S_URL — адрес API server; K3S_TOKEN — токен из /var/lib/rancher/k3s/server/node-token; "
        "при наличии K3S_URL установщик поднимает agent, а не server.",
    )

    sty.add_heading3(doc, "8.3. kubeconfig на локальном ПК", "m2_s08_3_kubeconfig")
    add_os_triple(
        doc,
        # win
        f"mkdir -p /c/Users/$USER/.kube\n"
        f"scp -i /c/Users/$USER/.ssh/id_ed25519 root@{K8S_MASTER_IP}:/etc/rancher/k3s/k3s.yaml /c/Users/$USER/.kube/selfhosted-greeting.yaml\n"
        f"sed -i 's/127.0.0.1/{K8S_MASTER_IP}/' /c/Users/$USER/.kube/selfhosted-greeting.yaml\n"
        "export KUBECONFIG=/c/Users/$USER/.kube/selfhosted-greeting.yaml\n"
        "kubectl get nodes",
        # mac
        f"mkdir -p ~/.kube\n"
        f"scp -i ~/.ssh/id_ed25519 root@{K8S_MASTER_IP}:/etc/rancher/k3s/k3s.yaml ~/.kube/selfhosted-greeting.yaml\n"
        f"sed -i '' 's/127.0.0.1/{K8S_MASTER_IP}/' ~/.kube/selfhosted-greeting.yaml\n"
        "export KUBECONFIG=~/.kube/selfhosted-greeting.yaml\n"
        "kubectl get nodes",
        # linux
        f"mkdir -p ~/.kube\n"
        f"scp -i ~/.ssh/id_ed25519 root@{K8S_MASTER_IP}:/etc/rancher/k3s/k3s.yaml ~/.kube/selfhosted-greeting.yaml\n"
        f"sed -i 's/127.0.0.1/{K8S_MASTER_IP}/' ~/.kube/selfhosted-greeting.yaml\n"
        "export KUBECONFIG=~/.kube/selfhosted-greeting.yaml\n"
        "kubectl get nodes",
    )
    sty.add_normal(
        doc,
        "Важно: на macOS у sed синтаксис sed -i '' 's/.../.../' (пустой суффикс бэкапа). "
        "На Linux/Git Bash — sed -i 's/.../.../'.",
    )

    sty.add_heading3(doc, "8.4. Разбор команд kubectl", "m2_s08_4_kubectl")
    sty.add_command_block(
        doc,
        "kubectl get nodes -o wide",
        "get — список ресурсов; nodes — тип; -o wide — дополнительные колонки (INTERNAL-IP, OS).",
    )
    sty.add_command_block(
        doc,
        "kubectl create namespace dev",
        "create namespace — создать пространство имён; dev — имя окружения.",
    )
    sty.add_command_block(
        doc,
        "kubectl -n dev get pods -o wide",
        "-n / --namespace — область; get pods — поды; -o wide — IP и NODE.",
    )
    sty.add_command_block(
        doc,
        "kubectl -n dev describe pod <POD>",
        "describe — события и состояние; полезно при ImagePullBackOff / CrashLoopBackOff.",
    )
    sty.add_command_block(
        doc,
        "kubectl -n dev logs deploy/greeting-service --tail=100",
        "logs — stdout контейнера; deploy/NAME — выбрать под Deployment; --tail — последние N строк.",
    )
    sty.add_command_block(
        doc,
        "kubectl -n dev expose deployment greeting-service --type=NodePort --name=greeting-service --port=80 --target-port=8080",
        "--type=NodePort — опубликовать на порту каждой worker-ноды; --port — порт Service; "
        "--target-port — порт контейнера приложения.",
    )
    sty.add_citation(
        doc,
        "https://kubernetes.io/docs/reference/kubectl/",
        "kubectl controls the Kubernetes cluster manager.",
        "kubectl управляет менеджером кластера Kubernetes.",
    )

    # ── 9 Traefik ────────────────────────────────────────────────────────
    sty.add_heading2(doc, "9. Кластер Traefik на отдельных VPS", "m2_s09_traefik")
    sty.add_normal(
        doc,
        "Traefik edge — два VPS (traefik-1, traefik-2) с Docker. Floating IP указывает на активную ноду. "
        "Это НЕ Ingress Controller внутри k3s и НЕ услуга облака.",
    )

    sty.add_heading3(doc, "9.1. Установка Traefik", "m2_s09_1_install")
    sty.add_platform_block(
        doc,
        "На traefik-1 и traefik-2 (после SSH)",
        "mkdir -p /opt/traefik/letsencrypt /opt/traefik/dynamic\n"
        "cat > /opt/traefik/docker-compose.yml << 'EOF'\n"
        "services:\n"
        "  traefik:\n"
        "    image: traefik:v3.1\n"
        "    restart: always\n"
        "    ports:\n"
        "      - \"80:80\"\n"
        "      - \"443:443\"\n"
        "      - \"8080:8080\"\n"
        "    volumes:\n"
        "      - /var/run/docker.sock:/var/run/docker.sock:ro\n"
        "      - ./traefik.yml:/etc/traefik/traefik.yml:ro\n"
        "      - ./dynamic:/etc/traefik/dynamic:ro\n"
        "      - ./letsencrypt:/letsencrypt\n"
        "EOF\n"
        "cat > /opt/traefik/traefik.yml << 'EOF'\n"
        "api:\n"
        "  dashboard: true\n"
        "  insecure: true\n"
        "entryPoints:\n"
        "  web:\n"
        "    address: \":80\"\n"
        "providers:\n"
        "  file:\n"
        "    directory: /etc/traefik/dynamic\n"
        "    watch: true\n"
        "EOF\n"
        "cd /opt/traefik && docker compose up -d\n"
        "docker ps",
    )
    sty.add_normal(
        doc,
        "Ключи: image traefik:v3.1 — официальный образ; ports 80/443 — вход из интернета; "
        "providers.file — маршруты из YAML без Docker labels приложения; "
        "docker compose up -d — фоновый запуск.",
    )

    sty.add_heading3(doc, "9.2. Маршрут к приложению", "m2_s09_2_route")
    sty.add_normal(
        doc,
        "Сначала узнайте NodePort сервиса в k3s, затем создайте dynamic-конфиг Traefik "
        f"с Host(`{DOMAIN}`) на IP worker + NodePort.",
    )
    sty.add_platform_block(
        doc,
        "Локальный ПК — любой ОС (kubectl)",
        "export KUBECONFIG=~/.kube/selfhosted-greeting.yaml   # на Windows Git Bash: /c/Users/$USER/.kube/...\n"
        "kubectl -n dev get svc greeting-service -o wide\n"
        "# найдите PORT(S), например 80:30080/TCP → NodePort=30080",
    )
    sty.add_platform_block(
        doc,
        "На traefik-1 (и скопируйте тот же файл на traefik-2)",
        f"cat > /opt/traefik/dynamic/greeting.yml << 'EOF'\n"
        "http:\n"
        "  routers:\n"
        "    greeting:\n"
        f"      rule: \"Host(`{DOMAIN}`)\"\n"
        "      entryPoints:\n"
        "        - web\n"
        "      service: greeting-svc\n"
        "  services:\n"
        "    greeting-svc:\n"
        "      loadBalancer:\n"
        "        servers:\n"
        f"          - url: \"http://{K8S_WORKER_1_IP}:30080\"\n"
        "EOF\n"
        "# Traefik подхватит файл автоматически (watch: true)",
        yaml_block=True,
    )

    # ── 10 Registry ──────────────────────────────────────────────────────
    sty.add_heading2(doc, "10. Docker Registry на devtools", "m2_s10_registry")
    sty.add_platform_block(
        doc,
        "На devtools (после SSH)",
        "apt-get install -y apache2-utils\n"
        "mkdir -p /opt/registry/{data,auth}\n"
        "htpasswd -Bbn docker docker > /opt/registry/auth/htpasswd\n"
        "docker run -d --name registry --restart=always -p 5000:5000 \\\n"
        "  -v /opt/registry/data:/var/lib/registry \\\n"
        "  -v /opt/registry/auth:/auth \\\n"
        "  -e REGISTRY_AUTH=htpasswd \\\n"
        "  -e REGISTRY_AUTH_HTPASSWD_REALM=\"Registry Realm\" \\\n"
        "  -e REGISTRY_AUTH_HTPASSWD_PATH=/auth/htpasswd \\\n"
        "  registry:2\n"
        "ufw allow 5000/tcp\n"
        "curl -u docker:docker http://127.0.0.1:5000/v2/",
    )
    sty.add_normal(
        doc,
        "htpasswd -Bbn — bcrypt, batch, username password; -p 5000:5000 — публикация порта; "
        "REGISTRY_AUTH=htpasswd — базовая аутентификация.",
    )
    sty.add_citation(
        doc,
        "https://distribution.github.io/distribution/",
        "The Registry is a stateless, highly scalable server side application that stores and lets you distribute Docker images.",
        "Registry — серверное приложение без состояния, которое хранит и позволяет распространять Docker-образы.",
    )
    sty.add_normal(doc, "На worker-нодах k3s разрешите insecure registry (HTTP), если нет TLS:")
    sty.add_platform_block(
        doc,
        "На каждой k8s-worker (и master, если тянет образы)",
        f"mkdir -p /etc/rancher/k3s\n"
        "cat > /etc/rancher/k3s/registries.yaml << EOF\n"
        "mirrors:\n"
        f"  \"{DEVTOOLS_IP}:5000\":\n"
        "    endpoint:\n"
        f"      - \"http://{DEVTOOLS_IP}:5000\"\n"
        "configs:\n"
        f"  \"{DEVTOOLS_IP}:5000\":\n"
        "    auth:\n"
        "      username: docker\n"
        "      password: docker\n"
        "EOF\n"
        "systemctl restart k3s-agent || systemctl restart k3s",
        yaml_block=True,
    )

    # ── 11 GitLab ────────────────────────────────────────────────────────
    sty.add_heading2(doc, "11. GitLab CE на devtools", "m2_s11_gitlab")
    sty.add_platform_block(
        doc,
        "На devtools (после SSH)",
        "apt-get install -y curl openssh-server ca-certificates tzdata perl\n"
        "curl -fsSL https://packages.gitlab.com/install/repositories/gitlab/gitlab-ce/script.deb.sh | bash\n"
        f"EXTERNAL_URL=\"http://{DEVTOOLS_IP}\" apt-get install -y gitlab-ce\n"
        "gitlab-ctl status\n"
        "cat /etc/gitlab/initial_root_password",
    )
    sty.add_normal(
        doc,
        "EXTERNAL_URL — публичный URL GitLab; gitlab-ctl status — состояние сервисов; "
        "initial_root_password — пароль root (удаляется через 24 часа).",
    )
    sty.add_normal(doc, "В UI: создать группу greeting-group и проект greeting-service (пустой, без README).")

    # ── 12 Runner ────────────────────────────────────────────────────────
    sty.add_heading2(doc, "12. GitLab Runner (self-hosted)", "m2_s12_runner")
    sty.add_normal(doc, "В GitLab UI: Project → Settings → CI/CD → Runners → New project runner → тег self-hosted.")
    sty.add_platform_block(
        doc,
        "На devtools (после SSH)",
        "curl -fsSL https://packages.gitlab.com/install/repositories/runner/gitlab-runner/script.deb.sh | bash\n"
        "apt-get install -y gitlab-runner\n"
        "usermod -aG docker gitlab-runner\n"
        f"gitlab-runner register --non-interactive \\\n"
        f"  --url \"http://{DEVTOOLS_IP}/\" \\\n"
        "  --token \"<runner-authentication-token>\" \\\n"
        "  --executor \"shell\" \\\n"
        "  --description \"devtools-runner\" \\\n"
        "  --tag-list \"self-hosted\"\n"
        "gitlab-runner verify",
    )
    sty.add_normal(
        doc,
        "--executor shell — команды на хосте (нужны docker/gradle/helm/kubectl на devtools); "
        "--tag-list — тег для jobs в .gitlab-ci.yml; usermod -aG docker — доступ Runner к Docker socket.",
    )
    sty.add_citation(
        doc,
        "https://docs.gitlab.com/runner/register/",
        "Registering a runner connects the runner with GitLab so it can receive jobs from GitLab CI/CD.",
        "Регистрация Runner связывает его с GitLab, чтобы он мог получать jobs из GitLab CI/CD.",
    )
    sty.add_normal(doc, "На devtools установите JDK, Helm и kubectl для shell-executor (один раз).")
    sty.add_platform_block(
        doc,
        "На devtools",
        "apt-get install -y openjdk-21-jdk-headless\n"
        "curl https://raw.githubusercontent.com/helm/helm/main/scripts/get-helm-3 | bash\n"
        "curl -LO \"https://dl.k8s.io/release/$(curl -L -s https://dl.k8s.io/release/stable.txt)/bin/linux/amd64/kubectl\"\n"
        "install -m 0755 kubectl /usr/local/bin/kubectl\n"
        "# скопируйте kubeconfig для пользователя gitlab-runner:\n"
        "mkdir -p /home/gitlab-runner/.kube\n"
        "cp /root/.kube/selfhosted-greeting.yaml /home/gitlab-runner/.kube/config\n"
        "chown -R gitlab-runner:gitlab-runner /home/gitlab-runner/.kube",
    )

    # ── 13 S3 ────────────────────────────────────────────────────────────
    sty.add_heading2(doc, "13. S3 Bucket (MinIO)", "m2_s13_s3")
    sty.add_normal(doc, "Вместо облачного twc_s3_bucket поднимаете MinIO на VPS storage.")
    sty.add_platform_block(
        doc,
        "На storage (после SSH)",
        "mkdir -p /opt/minio/data\n"
        "docker run -d --name minio --restart=always \\\n"
        "  -p 9000:9000 -p 9001:9001 \\\n"
        "  -e MINIO_ROOT_USER=minioadmin \\\n"
        "  -e MINIO_ROOT_PASSWORD='ChangeMe_StrongPass!' \\\n"
        "  -v /opt/minio/data:/data \\\n"
        "  quay.io/minio/minio server /data --console-address \":9001\"\n"
        "ufw allow 9000/tcp\n"
        "ufw allow 9001/tcp",
    )
    sty.add_normal(doc, "Создать bucket через mc (MinIO Client) с локального ПК или со storage:")
    add_os_triple(
        doc,
        # win — через docker mc
        f"docker run --rm -it --entrypoint=/bin/sh minio/mc -c \"\n"
        f"  mc alias set local http://{STORAGE_IP}:9000 minioadmin 'ChangeMe_StrongPass!' &&\n"
        "  mc mb local/greeting-artifacts &&\n"
        "  mc ls local\n"
        "\"",
        f"curl -O https://dl.min.io/client/mc/release/darwin-amd64/mc\n"
        "chmod +x mc && sudo mv mc /usr/local/bin/\n"
        f"mc alias set local http://{STORAGE_IP}:9000 minioadmin 'ChangeMe_StrongPass!'\n"
        "mc mb local/greeting-artifacts\n"
        "mc ls local",
        f"curl -O https://dl.min.io/client/mc/release/linux-amd64/mc\n"
        "chmod +x mc && sudo mv mc /usr/local/bin/\n"
        f"mc alias set local http://{STORAGE_IP}:9000 minioadmin 'ChangeMe_StrongPass!'\n"
        "mc mb local/greeting-artifacts\n"
        "mc ls local",
    )
    sty.add_normal(
        doc,
        "mc alias set — сохранить endpoint и ключи; mc mb — make bucket; "
        "S3_ENDPOINT / S3_ACCESS_KEY / S3_SECRET_KEY положите в GitLab CI Variables (Masked).",
    )

    # ── 14 PostgreSQL ────────────────────────────────────────────────────
    sty.add_heading2(doc, "14. PostgreSQL на отдельных VPS (Patroni)", "m2_s14_pg")
    sty.add_normal(
        doc,
        "Базу не ставим StatefulSet-ом в k3s. На схеме 1.1 это жёлтый блок: "
        "postgres-1 (primary) и postgres-2 (replica), связь с worker-нодами — частная сеть.",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, "На каждом postgres-VPS (после SSH, Ubuntu root):")
    sty.add_platform_block(
        doc,
        "На postgres-1 и postgres-2",
        "apt-get update\n"
        "apt-get install -y postgresql postgresql-contrib python3-pip python3-psycopg2 etcd\n"
        "systemctl stop postgresql\n"
        "systemctl disable postgresql\n"
        "# Patroni сам запускает Postgres; systemd postgresql не должен с ним спорить",
    )
    sty.add_normal(
        doc,
        "Patroni ставится официальным пакетом или через pip. Ниже — pip, как в документации Patroni "
        "(раздел Installation). Для production кластер etcd лучше из трёх узлов; для учебного стенда "
        "etcd может жить на postgres-1.",
    )
    sty.add_citation(
        doc,
        "https://patroni.readthedocs.io/en/latest/installation.html",
        "Patroni can be installed with pip: pip install patroni[dependencies] … "
        "Patroni packages may be available for your operating system … "
        "Once you have installed the PGDG repository for your OS you can install patroni.",
        "Patroni можно установить через pip: pip install patroni[зависимости]. "
        "Пакеты Patroni могут быть доступны для вашей ОС. После подключения репозитория PGDG "
        "Patroni ставят штатным пакетным менеджером (например apt-get install patroni).",
    )
    sty.add_platform_block(
        doc,
        "На postgres-1 и postgres-2",
        "pip3 install 'patroni[etcd]'",
    )
    sty.add_platform_block(
        doc,
        "Пример /etc/patroni.yml (имена и IP подставьте свои)",
        "scope: greeting-pg\n"
        "name: postgres-1\n"
        "restapi:\n"
        "  listen: 0.0.0.0:8008\n"
        f"  connect_address: {POSTGRES_PRIMARY_IP}:8008\n"
        "etcd:\n"
        f"  host: {POSTGRES_PRIMARY_IP}:2379\n"
        "bootstrap:\n"
        "  dcs:\n"
        "    ttl: 30\n"
        "    loop_wait: 10\n"
        "    postgresql:\n"
        "      use_pg_rewind: true\n"
        "  initdb:\n"
        "    - encoding: UTF8\n"
        "    - data-checksums\n"
        "postgresql:\n"
        "  listen: 0.0.0.0:5432\n"
        f"  connect_address: {POSTGRES_PRIMARY_IP}:5432\n"
        "  data_dir: /var/lib/postgresql/16/main\n"
        "  authentication:\n"
        "    superuser:\n"
        "      username: postgres\n"
        "      password: ChangeMe_DbPass\n"
        "    replication:\n"
        "      username: replicator\n"
        "      password: ChangeMe_ReplPass",
        yaml_block=True,
    )
    sty.add_normal(
        doc,
        "На replica тот же файл, но name: postgres-2 и connect_address с IP replica. "
        "Первый запуск Patroni на postgres-1 инициализирует primary; второй узел подхватит роль replica. "
        "Порт 5432 открывайте только в частной сети (ufw: allow from подсети worker-нод).",
    )
    sty.add_platform_block(
        doc,
        "Проверка с postgres-1",
        "patronictl -c /etc/patroni.yml list",
    )
    sty.add_normal(
        doc,
        f"JDBC из приложения (k3s workers, частная сеть): "
        f"jdbc:postgresql://{POSTGRES_PRIMARY_IP}:5432/greeting "
        "После появления VIP/HAProxy перед Patroni подставьте VIP вместо IP primary.",
    )

    # ── 15 Helm ──────────────────────────────────────────────────────────
    sty.add_heading2(doc, "15. Secrets и первый деплой Helm", "m2_s15_helm")
    sty.add_normal(
        doc,
        "Репозиторий уже содержит Helm chart (infra/helm или charts/greeting-service — "
        "используйте путь из вашего проекта). Ниже — общий алгоритм.",
    )
    sty.add_platform_block(
        doc,
        "Локальный ПК — Windows (Git Bash) / macOS / Ubuntu",
        "export KUBECONFIG=~/.kube/selfhosted-greeting.yaml\n"
        "kubectl create namespace dev\n"
        "kubectl -n dev create secret generic greeting-service-secret \\\n"
        f"  --from-literal=DB_URL='jdbc:postgresql://{POSTGRES_PRIMARY_IP}:5432/greeting' \\\n"
        "  --from-literal=DB_USER=postgres \\\n"
        "  --from-literal=DB_PASSWORD='ChangeMe_DbPass'\n"
        f"docker login {DEVTOOLS_IP}:5000 -u docker -p docker\n"
        "# соберите и запушьте образ (или дождитесь pipeline):\n"
        f"docker build -t {DEVTOOLS_IP}:5000/greeting-service:manual1 app/\n"
        f"docker push {DEVTOOLS_IP}:5000/greeting-service:manual1\n"
        "helm upgrade --install greeting-service ./infra/helm/greeting-service \\\n"
        "  --namespace dev \\\n"
        f"  --set image.repository={DEVTOOLS_IP}:5000/greeting-service \\\n"
        "  --set image.tag=manual1 \\\n"
        "  --atomic --wait\n"
        "kubectl -n dev get pods,svc",
    )
    sty.add_citation(
        doc,
        "https://helm.sh/docs/helm/helm_upgrade/",
        "This command upgrades a release to a new version of a chart.",
        "Команда обновляет release до новой версии chart.",
    )

    # ── 16 DNS ───────────────────────────────────────────────────────────
    sty.add_heading2(doc, "16. DNS", "m2_s16_dns")
    sty.add_normal(
        doc,
        "DNS настраивается у регистратора домена (Dynadot и т.п.), не через Terraform. "
        "A-запись должна указывать на Floating IP Traefik, а не на IP worker k3s напрямую.",
    )
    sty.add_empty(doc)
    sty.add_normal(doc, f"Тип: A", consolas=True)
    sty.add_normal(doc, f"Host/Name: greeting-dev (или @ / полный FQDN — как требует панель)", consolas=True)
    sty.add_normal(doc, f"Value: {TRAEFIK_FLOATING_IP}", consolas=True)
    sty.add_normal(doc, "TTL: 300", consolas=True)
    sty.add_empty(doc)
    sty.add_normal(doc, "Проверка до распространения DNS (через Host-заголовок):")
    add_os_triple(
        doc,
        f"curl -v -H \"Host: {DOMAIN}\" http://{TRAEFIK_FLOATING_IP}/api/greeting",
        f"curl -v -H \"Host: {DOMAIN}\" http://{TRAEFIK_FLOATING_IP}/api/greeting",
        f"curl -v -H \"Host: {DOMAIN}\" http://{TRAEFIK_FLOATING_IP}/api/greeting",
    )
    sty.add_normal(doc, "После TTL:")
    sty.add_platform_block(
        doc,
        "Любая ОС",
        f"nslookup {DOMAIN}\n"
        f"curl -v http://{DOMAIN}/api/greeting",
    )

    # ── 17 Pipeline ──────────────────────────────────────────────────────
    sty.add_heading2(doc, "17. GitLab CI/CD и подключение существующего репозитория", "m2_s17_pipeline")
    sty.add_normal(
        doc,
        "Код уже есть локально / в текущем git. Нужно добавить remote на ваш новый GitLab CE "
        "и запушить ветку вместе с ci/.gitlab-ci.yml.",
    )

    sty.add_heading3(doc, "17.1. Remote GitLab и push", "m2_s17_1_remote")
    add_os_triple(
        doc,
        f"cd '{REPO_WIN}'\n"
        "git remote -v\n"
        f"git remote add gitlab http://{DEVTOOLS_IP}/{GITLAB_GROUP}/{GITLAB_PROJECT}.git 2>/dev/null || true\n"
        "git push -u gitlab develop",
        f"cd {REPO_MAC}\n"
        "git remote -v\n"
        f"git remote add gitlab http://{DEVTOOLS_IP}/{GITLAB_GROUP}/{GITLAB_PROJECT}.git 2>/dev/null || true\n"
        "git push -u gitlab develop",
        f"cd {REPO_LIN}\n"
        "git remote -v\n"
        f"git remote add gitlab http://{DEVTOOLS_IP}/{GITLAB_GROUP}/{GITLAB_PROJECT}.git 2>/dev/null || true\n"
        "git push -u gitlab develop",
    )
    sty.add_normal(
        doc,
        "При HTTP push: Username = root (или ваш пользователь), Password = Personal Access Token "
        "со scope write_repository (GitLab → Preferences → Access Tokens).",
    )

    sty.add_heading3(doc, "17.2. CI/CD Variables", "m2_s17_2_vars")
    sty.add_normal(doc, "Project → Settings → CI/CD → Variables (Masked для секретов):")
    sty.add_platform_block(
        doc,
        "GitLab UI → Variables",
        f"REGISTRY_HOST         {DEVTOOLS_IP}:5000\n"
        "REGISTRY_USER         docker\n"
        "REGISTRY_PASSWORD     docker\n"
        "IMAGE_NAME            greeting-service\n"
        "HELM_RELEASE_NAME     greeting-service\n"
        f"S3_ENDPOINT           http://{STORAGE_IP}:9000\n"
        "S3_ACCESS_KEY         minioadmin\n"
        "S3_SECRET_KEY         ChangeMe_StrongPass!\n"
        "S3_BUCKET             greeting-artifacts",
    )

    sty.add_heading3(doc, "17.3. Пример ci/.gitlab-ci.yml", "m2_s17_3_yml")
    sty.add_platform_block(
        doc,
        "Файл в репозитории: ci/.gitlab-ci.yml",
        "stages:\n"
        "  - build\n"
        "  - docker\n"
        "  - deploy\n"
        "\n"
        "default:\n"
        "  tags:\n"
        "    - self-hosted\n"
        "\n"
        "build-and-test:\n"
        "  stage: build\n"
        "  script:\n"
        "    - ./gradlew clean test bootJar --no-daemon\n"
        "  artifacts:\n"
        "    paths:\n"
        "      - app/build/libs/*.jar\n"
        "\n"
        "build-and-push-docker:\n"
        "  stage: docker\n"
        "  script:\n"
        "    - echo \"$REGISTRY_PASSWORD\" | docker login -u \"$REGISTRY_USER\" --password-stdin \"$REGISTRY_HOST\"\n"
        "    - docker build -t \"$REGISTRY_HOST/$IMAGE_NAME:$CI_COMMIT_SHORT_SHA\" app/\n"
        "    - docker push \"$REGISTRY_HOST/$IMAGE_NAME:$CI_COMMIT_SHORT_SHA\"\n"
        "\n"
        "deploy-dev:\n"
        "  stage: deploy\n"
        "  script:\n"
        "    - helm upgrade --install \"$HELM_RELEASE_NAME\" ./infra/helm/greeting-service\n"
        "        --namespace dev --create-namespace\n"
        "        --set image.repository=\"$REGISTRY_HOST/$IMAGE_NAME\"\n"
        "        --set image.tag=\"$CI_COMMIT_SHORT_SHA\"\n"
        "        --atomic --wait\n"
        "  only:\n"
        "    - develop\n"
        "\n"
        "deploy-prod:\n"
        "  stage: deploy\n"
        "  when: manual\n"
        "  script:\n"
        "    - helm upgrade --install \"$HELM_RELEASE_NAME\" ./infra/helm/greeting-service\n"
        "        --namespace prod --create-namespace\n"
        "        --set image.repository=\"$REGISTRY_HOST/$IMAGE_NAME\"\n"
        "        --set image.tag=\"$CI_COMMIT_SHORT_SHA\"\n"
        "        --atomic --wait\n"
        "  only:\n"
        "    - main",
        yaml_block=True,
    )
    sty.add_normal(
        doc,
        "Если CI config path в GitLab не корневой .gitlab-ci.yml, укажите ci/.gitlab-ci.yml "
        "в Settings → CI/CD → General pipelines → CI/CD configuration file.",
    )
    sty.add_citation(
        doc,
        "https://docs.gitlab.com/ee/ci/pipelines/settings.html#specify-a-custom-cicd-configuration-file",
        "You can specify an alternate filename or path for the CI/CD configuration file.",
        "Можно указать альтернативное имя файла или путь к файлу конфигурации CI/CD.",
    )

    # ── 18 verify ────────────────────────────────────────────────────────
    sty.add_heading2(doc, "18. Проверка результата", "m2_s18_verify")
    sty.add_normal(doc, "1. GitLab → Build → Pipelines — pipeline зелёный.")
    sty.add_normal(doc, "2. kubectl -n dev get pods — Running / Ready 1/1.")
    sty.add_normal(doc, "3. curl через Traefik Floating IP с Host-заголовком.")
    sty.add_normal(doc, f"4. curl http://{DOMAIN}/api/greeting — после DNS.")
    sty.add_normal(doc, "5. mc ls local/greeting-artifacts — bucket доступен.")

    # ── 19 errors ────────────────────────────────────────────────────────
    sty.add_heading2(doc, "19. Типичные ошибки", "m2_s19_errors")
    sty.add_normal(doc, "• ImagePullBackOff — insecure registry / неверный пароль / registries.yaml на worker.")
    sty.add_normal(doc, "• Runner offline — gitlab-runner status; тег self-hosted не совпал.")
    sty.add_normal(doc, "• 404 от Traefik — Host() не совпадает с DNS-именем; NodePort устарел.")
    sty.add_normal(doc, "• DNS ок, сайт нет — Floating IP не на traefik-1; ufw закрыл 80.")
    sty.add_normal(doc, "• helm atomic rollback — смотрите kubectl -n dev describe pod / events.")
    sty.add_normal(doc, "• sed на macOS — забыли sed -i ''.")

    # ── 20 summary ───────────────────────────────────────────────────────
    sty.add_heading2(doc, "20. Финальная сводка", "m2_s20_summary")
    sty.add_normal(doc, "Порядок первого запуска:")
    sty.add_normal(doc, "1. Заказать VPS + Floating IP у любого провайдера.")
    sty.add_normal(doc, "2. SSH (Win/macOS/Linux) → подготовка ОС + Docker.")
    sty.add_normal(doc, "3. k3s master/workers (--disable traefik) → kubeconfig.")
    sty.add_normal(doc, "4. Traefik-кластер на отдельных VPS → dynamic route.")
    sty.add_normal(doc, "5. Registry + GitLab CE + Runner на devtools.")
    sty.add_normal(doc, "6. MinIO bucket на storage.")
    sty.add_normal(doc, "7. PostgreSQL primary+replica на отдельных VPS (Patroni) + Secret с JDBC.")
    sty.add_normal(doc, "8. DNS A → Traefik Floating IP.")
    sty.add_normal(doc, "9. git remote add gitlab → push → pipeline → Helm → проверка curl.")
    sty.add_empty(doc)
    sty.add_normal(
        doc,
        "Итог: тот же функциональный контур, что в Части 2, но без Terraform, без Timeweb-managed "
        "сервисов и без Bitbucket — только CLI, root и переносимые VPS.",
    )


def main() -> None:
    sty.reset_bookmarks()
    ensure_scheme11()
    render_tech_map()
    render_architecture()
    render_architecture_traffic()
    render_cicd()
    missing = [p for p in (IMG_TECH, IMG_ARCH, IMG_TRAFFIC, IMG_CICD) if not p.exists()]
    if missing:
        raise FileNotFoundError(f"Diagrams missing: {missing}")

    template = resolve_template()
    shutil.copy2(template, BUILD)
    doc = Document(BUILD)
    sty.clear_document_body(doc)
    build_document(doc)
    doc.save(BUILD)
    shutil.copy2(BUILD, OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
